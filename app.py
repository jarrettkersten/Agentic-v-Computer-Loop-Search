"""
Loop Comparison App — Cora ADO Edition
Compares two search approaches for querying the CoraSystems/PPM Azure DevOps repository:
  1. Agentic Loop  — Claude iteratively drives search strategy via tool calls
                     (ado_code_search + ado_read_file), deciding WHAT to look for next
                     mimicking a human navigating the ADO web interface

ALL answers come EXCLUSIVELY from the Cora PPM ADO repository.
No external sources. No general knowledge fallback.
"""

import os
import sys
import json
import re
import asyncio
import base64
import io
import csv
import time
import threading
import uuid
import queue
import secrets
import functools
import urllib.parse
import hashlib
import requests
from datetime import datetime, timedelta
from flask import Flask, render_template, request, jsonify, Response, stream_with_context, redirect, url_for, session, g
import anthropic
from dotenv import load_dotenv

# ---------------------------------------------------------------------------
# Bootstrap
# ---------------------------------------------------------------------------
load_dotenv()

# Playwright on Windows requires ProactorEventLoop
if sys.platform == "win32":
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
MODEL             = "claude-sonnet-4-6"
ALLOWED_MODELS    = {
    "claude-sonnet-4-6",
    "claude-opus-4-6",
    "claude-haiku-4-5-20251001",
}

def _resolve_model(data: dict | None) -> str:
    """Return a validated model string from a request payload, falling back to MODEL."""
    if data:
        req_model = (data.get("model") or "").strip()
        if req_model in ALLOWED_MODELS:
            return req_model
    return MODEL

ADO_ORG           = "CoraSystems"
ADO_PROJECT       = "PPM"
ADO_REPO          = "ppm"
ADO_BASE_URL      = f"https://dev.azure.com/{ADO_ORG}"
ADO_SEARCH_URL    = (
    f"https://almsearch.dev.azure.com/{ADO_ORG}/{ADO_PROJECT}"
    f"/_apis/search/codesearchresults?api-version=7.0"
)

# Cost estimation — USD per million tokens, per model
MODEL_PRICING = {
    "claude-sonnet-4-6":          {"input": 3.00,  "output": 15.00},
    "claude-opus-4-6":            {"input": 5.00,  "output": 25.00},
    "claude-haiku-4-5-20251001":  {"input": 1.00,  "output": 5.00},
}
# Legacy env-var overrides apply to the default model only
COST_PER_M_INPUT  = float(os.environ.get("COST_PER_M_INPUT",  "3.00"))
COST_PER_M_OUTPUT = float(os.environ.get("COST_PER_M_OUTPUT", "15.00"))
MODEL_PRICING[MODEL] = {"input": COST_PER_M_INPUT, "output": COST_PER_M_OUTPUT}

def _model_costs(model: str | None = None) -> tuple[float, float]:
    """Return (cost_per_m_input, cost_per_m_output) for the given model."""
    p = MODEL_PRICING.get(model or MODEL, MODEL_PRICING[MODEL])
    return p["input"], p["output"]

MAX_AGENTIC_ITERATIONS = 12    # more headroom for complex multi-file questions
MAX_SEARCH_RESULTS     = 12    # broader initial candidate pool per search term
MAX_FILE_CHARS         = 15000 # read more of each file — key methods are often deep
MAX_FILES_TO_READ      = 18    # read more files before synthesising

# Time-based search control ─────────────────────────────────────────────────
# Agentic loop pauses at these elapsed-second marks and asks the user whether
# to continue. Total hard stop is at SEARCH_HARD_LIMIT_SEC.
SEARCH_PAUSE_FIRST_SEC  = 300   # 5 min  — first "continue?" prompt
SEARCH_PAUSE_REPEAT_SEC = 120   # 2 min  — subsequent prompts
SEARCH_HARD_LIMIT_SEC   = 660   # 11 min — hard stop regardless

# ---------------------------------------------------------------------------
# Detect optional libraries at import time so we give clear errors if missing
try:
    import msal as _msal_mod
    MSAL_AVAILABLE = True
except ImportError:
    MSAL_AVAILABLE = False

# ---------------------------------------------------------------------------
# Entra ID (Azure AD) authentication config
# ---------------------------------------------------------------------------
ENTRA_CLIENT_ID     = os.environ.get("ENTRA_CLIENT_ID", "")
ENTRA_CLIENT_SECRET = os.environ.get("ENTRA_CLIENT_SECRET", "")
ENTRA_TENANT_ID     = os.environ.get("ENTRA_TENANT_ID", "")
ENTRA_REDIRECT_URI  = os.environ.get("ENTRA_REDIRECT_URI", "")
ENTRA_AUTHORITY      = f"https://login.microsoftonline.com/{ENTRA_TENANT_ID}" if ENTRA_TENANT_ID else ""
ENTRA_SCOPES         = ["User.Read"]

# Toggle auth on/off (disabled if no client ID is configured)
AUTH_DISABLED = not ENTRA_CLIENT_ID

# Comma-separated list of admin email addresses
ADMIN_EMAILS = [e.strip().lower() for e in os.environ.get("ADMIN_EMAILS", "").split(",") if e.strip()]

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", secrets.token_hex(32))

# Custom JSON serialisation so PostgreSQL datetime/date objects work with jsonify
from flask.json.provider import DefaultJSONProvider
import decimal

class _CustomJSONProvider(DefaultJSONProvider):
    def default(self, o):
        if isinstance(o, (datetime,)):
            return o.isoformat()
        if isinstance(o, decimal.Decimal):
            return float(o)
        try:
            from datetime import date as _date_type
            if isinstance(o, _date_type):
                return o.isoformat()
        except Exception:
            pass
        return super().default(o)

app.json_provider_class = _CustomJSONProvider
app.json = _CustomJSONProvider(app)

# ---------------------------------------------------------------------------
# Agentic search job store — supports mid-search "continue?" prompts via SSE
# ---------------------------------------------------------------------------
# Each running agentic job gets an entry:
#   job_id -> {
#     "event_queue":  queue.Queue   — server pushes SSE events here
#     "continue_evt": threading.Event — frontend signals "continue" by setting this
#     "stop_evt":     threading.Event — frontend signals "stop" by setting this
#   }
_agentic_jobs: dict = {}
_agentic_jobs_lock  = threading.Lock()

def _job_create(job_id: str):
    with _agentic_jobs_lock:
        _agentic_jobs[job_id] = {
            "event_queue":  queue.Queue(),
            "continue_evt": threading.Event(),
            "stop_evt":     threading.Event(),
        }

def _job_get(job_id: str):
    with _agentic_jobs_lock:
        return _agentic_jobs.get(job_id)

def _job_remove(job_id: str):
    with _agentic_jobs_lock:
        _agentic_jobs.pop(job_id, None)

# ---------------------------------------------------------------------------
# Usage tracking — PostgreSQL
# ---------------------------------------------------------------------------
import psycopg2
import psycopg2.extras
from psycopg2 import pool as _pg_pool

DATABASE_URL = os.environ.get("DATABASE_URL", "")
if not DATABASE_URL:
    print("[db] WARNING: DATABASE_URL not set — database features will be unavailable", flush=True)
else:
    print(f"[db] DATABASE_URL is set (host hidden for security)", flush=True)

# Simple thread-safe connection pool (min 1, max 10 connections)
_db_pool = None
def _get_pool():
    global _db_pool
    if _db_pool is None and DATABASE_URL:
        _db_pool = _pg_pool.ThreadedConnectionPool(1, 10, DATABASE_URL)
    return _db_pool

def get_db_conn():
    """Get a connection from the pool."""
    p = _get_pool()
    if p is None:
        raise RuntimeError("DATABASE_URL is not configured")
    return p.getconn()

def put_db_conn(conn):
    """Return a connection to the pool."""
    p = _get_pool()
    if p is not None:
        p.putconn(conn)


def init_db():
    """Create usage tracking tables if they don't exist, and migrate columns."""
    if not DATABASE_URL:
        print("[DB] Skipping init — no DATABASE_URL configured")
        return
    conn = None
    try:
        conn = get_db_conn()
        cur  = conn.cursor()
        # Enable trigram similarity for similar-query matching
        try:
            cur.execute("CREATE EXTENSION IF NOT EXISTS pg_trgm")
        except Exception:
            pass  # extension may already exist or require superuser
        cur.execute("""
            CREATE TABLE IF NOT EXISTS search_events (
                id               SERIAL PRIMARY KEY,
                ran_at           TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                loop_type        TEXT    NOT NULL,
                query            TEXT,
                branch           TEXT,
                status           TEXT    NOT NULL DEFAULT 'success',
                duration_sec     DOUBLE PRECISION,
                files_read       INTEGER DEFAULT 0,
                iterations       INTEGER DEFAULT 0,
                searches         INTEGER DEFAULT 0,
                confidence_level TEXT,
                input_tokens     INTEGER DEFAULT 0,
                output_tokens    INTEGER DEFAULT 0,
                total_tokens     INTEGER DEFAULT 0,
                error_message    TEXT,
                user_email       TEXT    DEFAULT ''
            )
        """)
        # Migrate existing tables: add columns if they don't already exist
        _safe_add_columns(cur, "search_events", [
            ("input_tokens",          "INTEGER DEFAULT 0"),
            ("output_tokens",         "INTEGER DEFAULT 0"),
            ("total_tokens",          "INTEGER DEFAULT 0"),
            ("user_email",            "TEXT DEFAULT ''"),
            ("response_sections",     "TEXT DEFAULT ''"),
            ("job_id",                "TEXT DEFAULT ''"),
            ("screenshot_context",    "TEXT DEFAULT ''"),
            ("answer_text",           "TEXT DEFAULT ''"),
            ("screenshot_b64",        "TEXT DEFAULT ''"),
            ("screenshot_mime",       "TEXT DEFAULT ''"),
            ("model",                 "TEXT DEFAULT ''"),
        ])
        _safe_add_columns(cur, "flagged_queries", [
            ("flag_type",         "TEXT NOT NULL DEFAULT 'inaccurate'"),
            ("container",         "TEXT DEFAULT ''"),
            ("request_article",   "INTEGER DEFAULT 0"),
            ("branch",            "TEXT DEFAULT ''"),
            ("confidence_level",  "TEXT DEFAULT ''"),
            ("duration_sec",      "DOUBLE PRECISION DEFAULT 0"),
            ("input_tokens",      "INTEGER DEFAULT 0"),
            ("output_tokens",     "INTEGER DEFAULT 0"),
            ("total_tokens",      "INTEGER DEFAULT 0"),
            ("files_read",        "INTEGER DEFAULT 0"),
            ("iterations",        "INTEGER DEFAULT 0"),
            ("searches",          "INTEGER DEFAULT 0"),
            ("search_event_id",   "INTEGER"),
            ("section_ratings",   "JSONB DEFAULT '{}'::jsonb"),
        ])
        # API jobs table — async query processing
        cur.execute("""
            CREATE TABLE IF NOT EXISTS api_jobs (
                id           TEXT    PRIMARY KEY,
                status       TEXT    NOT NULL DEFAULT 'pending',
                query        TEXT    NOT NULL,
                branch       TEXT,
                user_email   TEXT    DEFAULT 'external_api',
                result_json  TEXT,
                error        TEXT,
                created_at   TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                completed_at TIMESTAMPTZ
            )
        """)
        _safe_add_columns(cur, "api_jobs", [
            ("model", "TEXT DEFAULT ''"),
        ])
        # API keys table — multi-user API key management
        cur.execute("""
            CREATE TABLE IF NOT EXISTS api_keys (
                id         TEXT PRIMARY KEY,
                key_hash   TEXT NOT NULL,
                key_prefix TEXT NOT NULL,
                label      TEXT NOT NULL DEFAULT '',
                created_by TEXT NOT NULL DEFAULT '',
                created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                last_used  TIMESTAMPTZ,
                is_active  BOOLEAN NOT NULL DEFAULT TRUE
            )
        """)
        # Flagged queries table — stores responses users flag as unhelpful/missing
        cur.execute("""
            CREATE TABLE IF NOT EXISTS flagged_queries (
                id                TEXT PRIMARY KEY,
                query             TEXT NOT NULL,
                loop_type         TEXT NOT NULL,
                answer            TEXT,
                explanation       TEXT,
                flagged_by        TEXT,
                flagged_at        TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                status            TEXT NOT NULL DEFAULT 'pending',
                reviewed_by       TEXT,
                reviewed_at       TIMESTAMPTZ,
                admin_notes       TEXT,
                flag_type         TEXT NOT NULL DEFAULT 'inaccurate',
                container         TEXT DEFAULT '',
                request_article   INTEGER DEFAULT 0,
                branch            TEXT DEFAULT '',
                confidence_level  TEXT DEFAULT '',
                duration_sec      DOUBLE PRECISION DEFAULT 0,
                input_tokens      INTEGER DEFAULT 0,
                output_tokens     INTEGER DEFAULT 0,
                total_tokens      INTEGER DEFAULT 0,
                files_read        INTEGER DEFAULT 0,
                iterations        INTEGER DEFAULT 0,
                searches          INTEGER DEFAULT 0
            )
        """)
        # Codebase index — learned file/feature mapping to accelerate future searches
        cur.execute("""
            CREATE TABLE IF NOT EXISTS codebase_index (
                id               SERIAL PRIMARY KEY,
                file_path        TEXT NOT NULL,
                file_name        TEXT NOT NULL DEFAULT '',
                file_role        TEXT NOT NULL DEFAULT '',
                feature_area     TEXT NOT NULL DEFAULT '',
                keywords         TEXT NOT NULL DEFAULT '',
                related_files    TEXT NOT NULL DEFAULT '',
                description      TEXT NOT NULL DEFAULT '',
                branch           TEXT NOT NULL DEFAULT '',
                discovered_at    TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                last_verified    TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                hit_count        INTEGER NOT NULL DEFAULT 1,
                source_query     TEXT NOT NULL DEFAULT ''
            )
        """)
        # Create index for fast keyword lookups
        cur.execute("SAVEPOINT idx_check")
        try:
            cur.execute("""
                CREATE INDEX IF NOT EXISTS idx_codebase_index_keywords
                ON codebase_index USING gin (keywords gin_trgm_ops)
            """)
            cur.execute("RELEASE SAVEPOINT idx_check")
        except Exception:
            cur.execute("ROLLBACK TO SAVEPOINT idx_check")
        cur.execute("SAVEPOINT idx_check2")
        try:
            cur.execute("""
                CREATE INDEX IF NOT EXISTS idx_codebase_index_feature
                ON codebase_index USING gin (feature_area gin_trgm_ops)
            """)
            cur.execute("RELEASE SAVEPOINT idx_check2")
        except Exception:
            cur.execute("ROLLBACK TO SAVEPOINT idx_check2")
        # Unique constraint on file_path + branch to prevent duplicates
        cur.execute("SAVEPOINT uniq_check")
        try:
            cur.execute("""
                CREATE UNIQUE INDEX IF NOT EXISTS idx_codebase_index_file_branch
                ON codebase_index (file_path, branch)
            """)
            cur.execute("RELEASE SAVEPOINT uniq_check")
        except Exception:
            cur.execute("ROLLBACK TO SAVEPOINT uniq_check")

        # ── Add enriched context columns to codebase_index (migration) ──
        for col_name, col_def in [
            ("module",     "TEXT NOT NULL DEFAULT ''"),
            ("layer",      "TEXT NOT NULL DEFAULT ''"),
            ("technology", "TEXT NOT NULL DEFAULT ''"),
            ("ado_url",    "TEXT NOT NULL DEFAULT ''"),
        ]:
            cur.execute("SAVEPOINT col_add")
            try:
                cur.execute(f"ALTER TABLE codebase_index ADD COLUMN {col_name} {col_def}")
                cur.execute("RELEASE SAVEPOINT col_add")
                print(f"[init_db] Added column codebase_index.{col_name}")
            except Exception:
                cur.execute("ROLLBACK TO SAVEPOINT col_add")
        # Add trigram index on description for similarity search
        cur.execute("SAVEPOINT idx_desc")
        try:
            cur.execute("""
                CREATE INDEX IF NOT EXISTS idx_codebase_index_description
                ON codebase_index USING gin (description gin_trgm_ops)
            """)
            cur.execute("RELEASE SAVEPOINT idx_desc")
        except Exception:
            cur.execute("ROLLBACK TO SAVEPOINT idx_desc")
        # Add trigram index on module for similarity search
        cur.execute("SAVEPOINT idx_mod")
        try:
            cur.execute("""
                CREATE INDEX IF NOT EXISTS idx_codebase_index_module
                ON codebase_index USING gin (module gin_trgm_ops)
            """)
            cur.execute("RELEASE SAVEPOINT idx_mod")
        except Exception:
            cur.execute("ROLLBACK TO SAVEPOINT idx_mod")

        # One-time data fix: backfill model for historical records
        # IDs 1-16 were Sonnet, IDs 17-19 were Opus
        cur.execute("""
            UPDATE search_events
            SET model = 'claude-sonnet-4-6'
            WHERE id BETWEEN 1 AND 16
              AND (model IS NULL OR model = '')
        """)
        cur.execute("""
            UPDATE search_events
            SET model = 'claude-opus-4-6'
            WHERE id IN (17, 18, 19)
              AND (model IS NULL OR model = '' OR model = 'claude-sonnet-4-6')
        """)

        conn.commit()

    except Exception as e:
        print(f"[DB] Warning: could not init tracking tables: {e}")
        if conn:
            conn.rollback()
    finally:
        if conn:
            put_db_conn(conn)


def _safe_add_columns(cur, table, columns):
    """Add columns to a table, ignoring errors if they already exist."""
    for col_name, col_type in columns:
        cur.execute("SAVEPOINT add_col")
        try:
            cur.execute(f"ALTER TABLE {table} ADD COLUMN {col_name} {col_type}")
            cur.execute("RELEASE SAVEPOINT add_col")
        except Exception:
            # Column already exists — rollback the failed statement and continue
            cur.execute("ROLLBACK TO SAVEPOINT add_col")


# ---------------------------------------------------------------------------
# Codebase Index — learned file/feature mappings
# ---------------------------------------------------------------------------

def bump_index_hit_counts(file_paths: list, branch: str):
    """Increment hit_count for files that were returned from the codebase index.

    Called after lookup_codebase_index returns results so we can track which
    indexed files are most frequently relevant to user queries.
    """
    if not DATABASE_URL or not file_paths:
        return
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("""
            UPDATE codebase_index
            SET hit_count = hit_count + 1,
                last_verified = NOW()
            WHERE file_path = ANY(%s) AND branch = %s
        """, (file_paths, branch))
        conn.commit()
        print(f"[Index] Bumped hit_count for {len(file_paths)} file(s)")
    except Exception as e:
        print(f"[Index] Warning: could not bump hit counts: {e}")
        if conn:
            conn.rollback()
    finally:
        if conn:
            put_db_conn(conn)


def lookup_codebase_index(query: str, branch: str, limit: int = 20) -> list:
    """Search the codebase index for files relevant to a query.

    Uses trigram similarity on keywords and feature_area columns.
    Returns a list of dicts sorted by relevance (hit_count * similarity).
    """
    if not DATABASE_URL:
        return []
    import re
    # Extract meaningful terms from the query
    stop_words = {'the', 'a', 'an', 'is', 'are', 'was', 'were', 'what', 'how', 'does', 'do',
                  'in', 'on', 'at', 'to', 'for', 'of', 'and', 'or', 'but', 'not', 'this',
                  'that', 'it', 'from', 'with', 'by', 'as', 'be', 'has', 'have', 'had',
                  'can', 'could', 'would', 'should', 'will', 'when', 'where', 'why', 'which',
                  'cora', 'ppm'}
    terms = [w for w in re.findall(r'\b\w+\b', query.lower()) if w not in stop_words and len(w) > 2]
    if not terms:
        return []

    search_text = " ".join(terms)
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT file_path, file_name, file_role, feature_area, keywords,
                   related_files, description, hit_count, last_verified,
                   module, layer, technology, ado_url,
                   GREATEST(
                       similarity(keywords, %(search)s),
                       similarity(feature_area, %(search)s),
                       similarity(file_name, %(search)s),
                       similarity(description, %(search)s) * 0.8,
                       similarity(module, %(search)s)
                   ) AS sim_score
            FROM codebase_index
            WHERE branch = %(branch)s
              AND (
                  similarity(keywords, %(search)s) > 0.08
                  OR similarity(feature_area, %(search)s) > 0.15
                  OR similarity(file_name, %(search)s) > 0.2
                  OR similarity(description, %(search)s) > 0.12
                  OR similarity(module, %(search)s) > 0.15
              )
            ORDER BY (GREATEST(hit_count, 1) * GREATEST(
                similarity(keywords, %(search)s),
                similarity(feature_area, %(search)s),
                similarity(file_name, %(search)s),
                similarity(description, %(search)s) * 0.8,
                similarity(module, %(search)s)
            )) DESC
            LIMIT %(limit)s
        """, {"search": search_text, "branch": branch, "limit": limit})
        results = [dict(r) for r in cur.fetchall()]
        # Bump hit counts for returned files (fire-and-forget in background)
        if results:
            hit_paths = [r["file_path"] for r in results]
            try:
                bump_index_hit_counts(hit_paths, branch)
            except Exception:
                pass  # Non-critical; don't break the search
        return results
    except Exception as e:
        print(f"[Index] Warning: lookup failed: {e}")
        return []
    finally:
        if conn:
            put_db_conn(conn)


def format_index_hints(index_hits: list) -> str:
    """Format index lookup results as a system prompt section for Claude.

    Returns a string to inject into the system prompt, or empty string if no hits.
    Each entry now includes the file's description and the past queries that
    led to it, giving Claude real context about *why* each file matters.
    """
    if not index_hits:
        return ""

    lines = [
        "\n═══ KNOWN FILE INDEX (from past searches) ═══",
        "The following files were relevant in previous searches for similar questions.",
        "Use these as STARTING POINTS — read these files first before searching broadly.\n",
    ]

    # Group by feature area
    by_feature: dict = {}
    for h in index_hits:
        area = h.get("feature_area") or "General"
        if area not in by_feature:
            by_feature[area] = []
        by_feature[area].append(h)

    for area, files in by_feature.items():
        lines.append(f"  [{area}]")
        for f in files:
            role = f" ({f['file_role']})" if f.get("file_role") else ""
            hits = f.get("hit_count", 1)
            lines.append(f"    • {f['file_path']}{role}  [seen {hits}x]")
            # Add description if available
            desc = (f.get("description") or "").strip()
            if desc:
                lines.append(f"      Purpose: {desc}")
            # Show module, layer, and technology context
            module = (f.get("module") or "").strip()
            layer = (f.get("layer") or "").strip()
            tech = (f.get("technology") or "").strip()
            context_parts = []
            if module:
                context_parts.append(f"Module: {module}")
            if layer:
                context_parts.append(f"Layer: {layer}")
            if tech:
                context_parts.append(f"Tech: {tech}")
            if context_parts:
                lines.append(f"      {' | '.join(context_parts)}")
            # Show ADO link if available
            ado_url = (f.get("ado_url") or "").strip()
            if ado_url:
                lines.append(f"      ADO: {ado_url}")
            # Show past queries that led to this file (pipe-delimited in keywords)
            kw = (f.get("keywords") or "").strip()
            if "|" in kw:
                past_queries = [q.strip() for q in kw.split("|") if q.strip()]
                if past_queries:
                    lines.append(f"      Past queries: {'; '.join(past_queries[-5:])}")
            elif kw:
                lines.append(f"      Keywords: {kw}")
        lines.append("")

    lines.append("These are hints, not guarantees — files may have moved or been renamed. Verify by reading them.")
    lines.append("STILL search broadly after checking these — there may be files not yet indexed.\n")

    return "\n".join(lines)


def _detect_response_sections(answer: str) -> str:
    """Return a comma-separated list of which response sections are present."""
    if not answer:
        return ""
    sections = []
    if re.search(r"## Overview", answer, re.IGNORECASE):
        sections.append("Overview")
    if re.search(r"## User Guide", answer, re.IGNORECASE):
        sections.append("User Guide")
    if re.search(r"## Technical Reference", answer, re.IGNORECASE):
        sections.append("Technical Reference")
    return ", ".join(sections) if sections else "Unstructured"


def _split_answer_sections(answer: str) -> dict:
    """Split the agentic answer into the three container sections."""
    if not answer:
        return {"overview": "", "user_guide": "", "technical_reference": ""}

    overview_match = re.search(
        r"## Overview\s*\n([\s\S]*?)(?=\n---\s*\n|\n## User Guide|$)", answer, re.IGNORECASE
    )
    guide_match = re.search(
        r"## User Guide\s*\n([\s\S]*?)(?=\n---\s*\n|\n## Technical Reference|$)", answer, re.IGNORECASE
    )
    tech_match = re.search(
        r"## Technical Reference\s*\n([\s\S]*?)$", answer, re.IGNORECASE
    )

    overview = overview_match.group(1).strip() if overview_match else ""
    user_guide = guide_match.group(1).strip() if guide_match else ""
    technical_reference = tech_match.group(1).strip() if tech_match else ""

    # Fallback: if no sections found, put everything in overview
    if not overview and not user_guide and not technical_reference:
        overview = answer.strip()

    return {
        "overview": overview,
        "user_guide": user_guide,
        "technical_reference": technical_reference,
    }


# ---------------------------------------------------------------------------
# Shared helpers: persist search jobs to api_jobs for all search paths
# ---------------------------------------------------------------------------

def _persist_api_job(job_id, query, branch, user_email):
    """Create an api_jobs row so the result is retrievable via GET /api/query/<job_id>."""
    if not DATABASE_URL:
        return
    conn = None
    try:
        conn = get_db_conn()
        conn.cursor().execute(
            "INSERT INTO api_jobs (id, status, query, branch, user_email) "
            "VALUES (%s, 'processing', %s, %s, %s)",
            (job_id, query, branch, user_email),
        )
        conn.commit()
    except Exception as e:
        print(f"[DB] Warning: could not persist api_job {job_id}: {e}")
        if conn:
            conn.rollback()
    finally:
        if conn:
            put_db_conn(conn)


def _complete_api_job(job_id, result, screenshot_context=""):
    """Write the completed result to api_jobs."""
    if not DATABASE_URL:
        return
    raw_answer = result.get("answer", "")
    sections = _split_answer_sections(raw_answer)
    payload_data = {
        "success": True,
        "answer": raw_answer,
        "sections": sections,
        "metadata": {
            "branch":            result.get("branch", ""),
            "files_read":        result.get("total_files_read", 0),
            "iterations":        result.get("total_iterations", 0),
            "searches":          result.get("total_searches", 0),
            "confidence":        result.get("confidence", {}).get("level", ""),
            "elapsed_seconds":   result.get("elapsed_seconds", 0),
            "input_tokens":      result.get("input_tokens", 0),
            "output_tokens":     result.get("output_tokens", 0),
            "response_sections": _detect_response_sections(raw_answer),
        },
    }
    if screenshot_context:
        payload_data["screenshot_context"] = screenshot_context
    payload = json.dumps(payload_data)
    conn = None
    try:
        conn = get_db_conn()
        conn.cursor().execute(
            "UPDATE api_jobs SET status='completed', result_json=%s, completed_at=NOW() WHERE id=%s",
            (payload, job_id),
        )
        conn.commit()
    except Exception as e:
        print(f"[DB] Warning: could not complete api_job {job_id}: {e}")
        if conn:
            conn.rollback()
    finally:
        if conn:
            put_db_conn(conn)


def _fail_api_job(job_id, error_msg, elapsed_sec=0):
    """Mark an api_job as failed."""
    if not DATABASE_URL:
        return
    payload = json.dumps({
        "success": False,
        "error": error_msg,
        "metadata": {"total_time_sec": round(elapsed_sec, 1)},
    })
    conn = None
    try:
        conn = get_db_conn()
        conn.cursor().execute(
            "UPDATE api_jobs SET status='failed', error=%s, result_json=%s, completed_at=NOW() WHERE id=%s",
            (error_msg, payload, job_id),
        )
        conn.commit()
    except Exception as e:
        print(f"[DB] Warning: could not fail api_job {job_id}: {e}")
        if conn:
            conn.rollback()
    finally:
        if conn:
            put_db_conn(conn)


def track_search_event(loop_type: str, query: str, branch: str, status: str,
                       duration_sec: float, files_read: int = 0,
                       iterations: int = 0, searches: int = 0,
                       confidence_level: str = "", error_message: str = "",
                       input_tokens: int = 0, output_tokens: int = 0,
                       user_email: str = "", answer: str = "",
                       job_id: str = "", screenshot_context: str = "",
                       screenshot_b64: str = "", screenshot_mime: str = "",
                       model: str = ""):
    """Record a single loop search run."""
    if not DATABASE_URL:
        return
    response_sections = _detect_response_sections(answer)
    conn = None
    try:
        conn = get_db_conn()
        cur  = conn.cursor()
        cur.execute(
            """INSERT INTO search_events
               (loop_type, query, branch, status, duration_sec, files_read,
                iterations, searches, confidence_level,
                input_tokens, output_tokens, total_tokens, error_message,
                user_email, response_sections, job_id, screenshot_context,
                answer_text, screenshot_b64, screenshot_mime, model)
               VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
               RETURNING id""",
            (loop_type, (query or "")[:2000], branch or "", status,
             round(duration_sec, 2), files_read, iterations, searches,
             confidence_level or "",
             input_tokens, output_tokens, input_tokens + output_tokens,
             error_message or "", user_email or "", response_sections,
             job_id or "", screenshot_context or "", answer or "",
             screenshot_b64 or "", screenshot_mime or "", model or MODEL),
        )
        row = cur.fetchone()
        conn.commit()
        return row[0] if row else None
    except Exception as e:
        print(f"[DB] Warning: could not track search event: {e}")
        if conn:
            conn.rollback()
        return None
    finally:
        if conn:
            put_db_conn(conn)



def query_db(sql: str, params=None) -> list:
    """Run a SELECT and return rows as list of dicts."""
    if not DATABASE_URL:
        return []
    conn = None
    try:
        conn = get_db_conn()
        cur  = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(sql, params)
        rows = [dict(r) for r in cur.fetchall()]
        return rows
    except Exception as e:
        print(f"[DB] Warning: query failed: {e}")
        return []
    finally:
        if conn:
            put_db_conn(conn)


# Initialise tables on startup
init_db()


# ---------------------------------------------------------------------------
# Entra ID auth helpers
# ---------------------------------------------------------------------------

def _build_msal_app():
    return _msal_mod.ConfidentialClientApplication(
        ENTRA_CLIENT_ID,
        authority=ENTRA_AUTHORITY,
        client_credential=ENTRA_CLIENT_SECRET,
    )


def _current_user_email():
    if AUTH_DISABLED:
        return "anonymous"
    return session.get("user", {}).get("mail", session.get("user", {}).get("userPrincipalName", ""))


def _is_admin():
    if AUTH_DISABLED:
        return True
    email = _current_user_email().lower()
    return email in ADMIN_EMAILS


def login_required(f):
    @functools.wraps(f)
    def decorated(*args, **kwargs):
        if AUTH_DISABLED:
            return f(*args, **kwargs)
        if "user" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    @functools.wraps(f)
    def decorated(*args, **kwargs):
        if AUTH_DISABLED:
            return f(*args, **kwargs)
        if "user" not in session:
            return redirect(url_for("login"))
        if not _is_admin():
            return jsonify({"error": "Admin access required"}), 403
        return f(*args, **kwargs)
    return decorated


# ---------------------------------------------------------------------------
# ADO helpers
# ---------------------------------------------------------------------------

def _ado_headers() -> dict:
    pat   = os.environ.get("AZURE_DEVOPS_PAT", "")
    token = base64.b64encode(f":{pat}".encode()).decode()
    return {
        "Authorization": f"Basic {token}",
        "Content-Type":  "application/json",
    }


def ado_get_branches() -> list:
    """Return the branches that are actually indexed by ADO Code Search.

    The Code Search facets API caps returned facet values at ~10 per call,
    ordered by result count.  Newly-added branches with fewer indexed hits
    may fall below that cap and not appear.

    Fix: query with several common code tokens and union all branch names
    across the responses.  Falls back to the Git Refs API if every facets
    call fails.

    Display order: supported_release-* newest-first, then everything else.
    """
    import re as _re

    def _sort_key(b):
        # supported_release-X.Y.Z  → bucket 0, sorted newest-first (descending)
        # anything else             → bucket 1, sorted alphabetically
        m = _re.search(r"(\d+)\.(\d+)\.(\d+)", b)
        if b.lower().startswith("supported_release") and m:
            return (0, -int(m.group(1)), -int(m.group(2)), -int(m.group(3)), b)
        return (1, 0, 0, 0, b)

    # ── Primary: union facets from several search terms ───────────────────────
    # Different terms hit different files/branches; unioning covers branches
    # that might be below the per-call facet cap for any single term.
    _PROBE_TERMS = ["class", "public", "return", "void", "function"]
    all_branches: set = set()
    any_facets_ok = False
    for term in _PROBE_TERMS:
        try:
            body = {
                "searchText": term,
                "$skip": 0,
                "$top": 1,
                "filters": {
                    "Project":    [ADO_PROJECT],
                    "Repository": [ADO_REPO],
                },
                "includeFacets": True,
            }
            resp = requests.post(ADO_SEARCH_URL, headers=_ado_headers(), json=body, timeout=15)
            if resp.ok:
                facets = resp.json().get("facets", {})
                for entry in facets.get("Branch", []):
                    if entry.get("name"):
                        all_branches.add(entry["name"])
                        any_facets_ok = True
        except Exception as e:
            print(f"[branches] Facets probe '{term}' failed: {e}")

    if all_branches:
        print(f"[branches] Facets collected {len(all_branches)} branches: {sorted(all_branches)}")
        return sorted(all_branches, key=_sort_key)

    if not any_facets_ok:
        print("[branches] All facets calls failed — falling back to Git Refs API")

    # ── Fallback: Git Refs API filtered to supported_release branches ─────────
    BRANCH_PATTERN = _re.compile(r"^supported_release-\d+\.\d+\.\d+$")
    url = (
        f"{ADO_BASE_URL}/{ADO_PROJECT}/_apis/git/repositories/{ADO_REPO}"
        f"/refs?filter=heads/supported_release&api-version=7.0"
    )
    resp = requests.get(url, headers=_ado_headers(), timeout=10)
    resp.raise_for_status()
    branches = []
    for ref in resp.json().get("value", []):
        name = ref.get("name", "")
        if name.startswith("refs/heads/"):
            short = name[len("refs/heads/"):]
            if BRANCH_PATTERN.match(short):
                branches.append(short)
    return sorted(branches, key=_sort_key)


def ado_code_search(query, branch, top=None):
    """Search the Cora PPM ADO codebase and return matching file metadata.

    Now that the admin has enabled specific branches for Code Search, we pass
    the branch filter directly so results come from exactly the branch the user
    selected — no more searching the default branch and re-reading files from
    a different branch.

    ADO Code Search API requires $skip/$top (dollar-sign prefix).
    """
    if top is None:
        top = MAX_SEARCH_RESULTS

    filters: dict = {
        "Project":    [ADO_PROJECT],
        "Repository": [ADO_REPO],
    }
    # Pass the branch filter so results come from the selected branch directly.
    # If for any reason the branch isn't indexed, ADO falls back to results
    # from any indexed branch — callers still read files using the explicit
    # branch, so correctness is preserved either way.
    if branch:
        filters["Branch"] = [branch]

    body = {
        "searchText": query,
        "$skip": 0,
        "$top": top,
        "filters": filters,
    }

    resp = requests.post(ADO_SEARCH_URL, headers=_ado_headers(), json=body, timeout=15)

    # Raise with full ADO error body so we can diagnose failures
    if not resp.ok:
        try:
            detail = resp.json()
        except Exception:
            detail = resp.text[:500]
        raise requests.HTTPError(
            f"{resp.status_code} {resp.reason} — ADO detail: {detail}",
            response=resp,
        )

    results = []
    for r in resp.json().get("results", [])[:top]:
        entry = {
            "file_path":  r.get("path", ""),
            "file_name":  r.get("fileName", ""),
            "repository": r.get("repository", {}).get("name", ""),
        }
        # Extract content snippets from ADO's hits array.
        # Each hit contains the matching lines of code at a specific char offset.
        # Including these means callers can see matching code from ANY file — even
        # large files that would be truncated during a full read — without an extra
        # API call.  We keep the top 3 snippets per file, trimming whitespace.
        snippets = []
        for hit in r.get("hits", []):
            text = hit.get("content", "").strip()
            offset = hit.get("charOffset", 0)
            if text and len(snippets) < 3:
                snippets.append({"offset": offset, "content": text})
        if snippets:
            entry["snippets"] = snippets
        results.append(entry)
    return results


def ado_read_file(file_path, branch, char_offset=0):
    """Read a file from the Cora PPM ADO repo.

    Returns the FULL file content starting from char_offset.  No truncation —
    the entire file (from the offset onward) is returned so the model never
    misses methods or classes defined deeper in the file.
    """
    url = (
        f"{ADO_BASE_URL}/{ADO_PROJECT}/_apis/git/repositories/{ADO_REPO}/items"
        f"?path={requests.utils.quote(file_path)}"
        f"&versionDescriptor.version={branch}"
        f"&versionDescriptor.versionType=branch"
        f"&$format=text"
        f"&api-version=7.0"
    )
    resp = requests.get(url, headers=_ado_headers(), timeout=15)
    if not resp.ok:
        try:
            detail = resp.json()
        except Exception:
            detail = resp.text[:500]
        raise requests.HTTPError(
            f"{resp.status_code} {resp.reason} — ADO detail: {detail}",
            response=resp,
        )
    full_content = resp.text
    total_len    = len(full_content)

    # Return everything from the requested offset onward — no chunking.
    content = full_content[char_offset:]

    if not content:
        return f"[No content at offset {char_offset:,} — file is {total_len:,} chars total.]"

    # Append a size footer so the model (and logs) know how much was returned.
    content += f"\n\n[Full file returned — {len(content):,} chars"
    if char_offset > 0:
        content += f" (from offset {char_offset:,})"
    content += f" of {total_len:,} total.]"

    return content


def ado_get_file_commits(file_path, branch, top=3):
    """Fetch the most recent commits that touched a specific file on a branch.

    Returns a list of dicts:
      [{"commit_id": "abc123", "short_id": "abc123",
        "author": "Jane Doe", "author_email": "jane@co.com",
        "date": "2026-03-15T14:30:00Z",
        "comment": "Fix EAC submit validation",
        "url": "https://dev.azure.com/..."}]
    """
    url = (
        f"{ADO_BASE_URL}/{ADO_PROJECT}/_apis/git/repositories/{ADO_REPO}/commits"
        f"?searchCriteria.itemPath={requests.utils.quote(file_path)}"
        f"&searchCriteria.itemVersion.version={branch}"
        f"&searchCriteria.itemVersion.versionType=branch"
        f"&$top={top}"
        f"&api-version=7.0"
    )
    try:
        resp = requests.get(url, headers=_ado_headers(), timeout=15)
        if not resp.ok:
            return []
        data = resp.json()
        commits = []
        for c in data.get("value", []):
            author = c.get("author", {})
            commit_id = c.get("commitId", "")
            commits.append({
                "commit_id":    commit_id,
                "short_id":     commit_id[:8],
                "author":       author.get("name", "Unknown"),
                "author_email": author.get("email", ""),
                "date":         author.get("date", ""),
                "comment":      (c.get("comment") or "").strip().split("\n")[0],  # first line only
                "url":          f"https://dev.azure.com/{ADO_ORG}/{ADO_PROJECT}/_git/{ADO_REPO}/commit/{commit_id}",
            })
        return commits
    except Exception:
        return []


def extract_code_sections(answer_text, accumulated_files):
    """Parse Claude's answer for code blocks with file references and find their line numbers.

    Looks for patterns like:
      File: /code/path/to/File.cs, lines 120–145
      ```csharp
      ... code ...
      ```
    Also handles `File: path/File.cs` without explicit line numbers by matching
    the snippet content against accumulated_files to determine line ranges.

    Returns a list of dicts:
      [{"file_path": "/code/Foo.cs", "start_line": 120, "end_line": 145,
        "snippet_preview": "first 80 chars of snippet"}]
    """
    import re

    sections = []
    seen = set()  # (file_path, start_line) dedup key

    # Pattern 1: Explicit "File: path, lines N–M" followed by a code block
    explicit_pat = re.compile(
        r'[Ff]ile:\s*[`]?([^\n`,]+\.\w{1,10})[`]?\s*,?\s*[Ll]ines?\s+(\d+)\s*[–\-—to]+\s*(\d+)',
        re.MULTILINE
    )
    for m in explicit_pat.finditer(answer_text):
        fp = m.group(1).strip()
        start = int(m.group(2))
        end = int(m.group(3))
        key = (fp, start)
        if key not in seen:
            seen.add(key)
            sections.append({
                "file_path": fp,
                "start_line": start,
                "end_line": end,
                "snippet_preview": "",
            })

    # Pattern 2: Inline method references like `MethodName()` in `File.ext` (line N)
    inline_pat = re.compile(
        r'`([^`]+)`\s+in\s+`([^`]+\.\w{1,10})`\s*\(line\s+(\d+)\)',
        re.MULTILINE
    )
    for m in inline_pat.finditer(answer_text):
        fp = m.group(2).strip()
        line_num = int(m.group(3))
        # Create a small window around the referenced line
        start = max(1, line_num - 5)
        end = line_num + 15
        key = (fp, start)
        if key not in seen:
            seen.add(key)
            sections.append({
                "file_path": fp,
                "start_line": start,
                "end_line": end,
                "snippet_preview": m.group(1).strip(),
            })

    # Pattern 3: Code blocks preceded by any file reference — match snippet against accumulated_files
    block_pat = re.compile(
        r'[Ff]ile:\s*[`]?([^\n`,:]+\.\w{1,10})[`]?\s*\n\s*```\w*\n(.*?)```',
        re.DOTALL
    )
    for m in block_pat.finditer(answer_text):
        fp = m.group(1).strip()
        snippet = m.group(2).strip()
        if not snippet or len(snippet) < 10:
            continue

        # Try to find line numbers by matching snippet against file content
        start_line, end_line = _match_snippet_to_lines(fp, snippet, accumulated_files)
        if start_line:
            key = (fp, start_line)
            if key not in seen:
                seen.add(key)
                sections.append({
                    "file_path": fp,
                    "start_line": start_line,
                    "end_line": end_line,
                    "snippet_preview": snippet[:80],
                })

    return sections


def _match_snippet_to_lines(file_path, snippet, accumulated_files):
    """Match a code snippet against a file's full content to find line numbers.

    Returns (start_line, end_line) or (None, None) if not found.
    """
    # Normalise the file_path — accumulated_files may use full or partial paths
    content = None
    for key, val in accumulated_files.items():
        if key == file_path or key.endswith(file_path) or file_path.endswith(key.lstrip("/")):
            content = val
            break

    if not content:
        return None, None

    # Normalise whitespace for comparison
    snippet_lines = [l.strip() for l in snippet.splitlines() if l.strip()]
    content_lines = content.splitlines()

    if not snippet_lines:
        return None, None

    # Search for the first line of the snippet in the file
    first_line = snippet_lines[0]
    for i, cline in enumerate(content_lines):
        if first_line in cline.strip():
            # Verify a few more lines match
            match_count = 1
            for j in range(1, min(len(snippet_lines), 5)):
                if i + j < len(content_lines) and snippet_lines[j] in content_lines[i + j].strip():
                    match_count += 1
            if match_count >= min(len(snippet_lines), 3) or match_count == len(snippet_lines):
                start = i + 1  # 1-based
                end = start + len(snippet_lines) - 1
                return start, end

    return None, None


def fetch_blame_changes(answer_text, accumulated_files, branch):
    """Post-processing: fetch blame-level recent commits for code sections Claude highlighted.

    Instead of file-level commits, this extracts the specific code sections
    (with line ranges) that Claude identified in its answer, then fetches
    commits that touched those specific lines.

    Returns a list of dicts sorted by commit date (newest first):
      [{"file_path": "/src/Foo.cs", "file_name": "Foo.cs",
        "start_line": 120, "end_line": 145,
        "snippet_preview": "first 80 chars...",
        "ado_section_url": "https://dev.azure.com/.../commit/...?path=...&line=120...",
        "commits": [<commit_dict>, ...]}]
    """
    sections = extract_code_sections(answer_text, accumulated_files)

    if not sections:
        return []

    from concurrent.futures import ThreadPoolExecutor, as_completed

    def _fetch_section(section):
        fp = section["file_path"]
        commits = ado_get_file_commits(fp, branch, top=3)
        # Enrich each commit URL with line range parameters
        for c in commits:
            c["url"] = _build_ado_line_url(
                c["commit_id"], fp, branch,
                section["start_line"], section["end_line"]
            )
        return section, commits

    results = []
    with ThreadPoolExecutor(max_workers=6) as pool:
        futures = [pool.submit(_fetch_section, s) for s in sections]
        for future in as_completed(futures):
            try:
                section, commits = future.result()
                if commits:
                    file_name = section["file_path"].rsplit("/", 1)[-1] if "/" in section["file_path"] else section["file_path"]
                    results.append({
                        "file_path":       section["file_path"],
                        "file_name":       file_name,
                        "start_line":      section["start_line"],
                        "end_line":        section["end_line"],
                        "snippet_preview": section.get("snippet_preview", ""),
                        "ado_section_url": _build_ado_browse_line_url(
                            section["file_path"], branch,
                            section["start_line"], section["end_line"]
                        ),
                        "commits":         commits,
                    })
            except Exception:
                pass

    # Sort by most recent commit date
    results.sort(
        key=lambda c: c["commits"][0]["date"] if c["commits"] else "",
        reverse=True,
    )
    return results


def _build_ado_line_url(commit_id, file_path, branch, start_line, end_line):
    """Build an ADO commit URL that highlights specific lines.

    Format: https://dev.azure.com/{org}/{project}/_git/{repo}/commit/{commitId}
            ?path={path}&version=GB{branch}&line={start}&lineEnd={end}
            &lineStartColumn=1&lineEndColumn=25&type=2&lineStyle=plain&_a=files
    """
    path_enc = requests.utils.quote(file_path)
    return (
        f"https://dev.azure.com/{ADO_ORG}/{ADO_PROJECT}/_git/{ADO_REPO}/commit/{commit_id}"
        f"?path={path_enc}&version=GB{requests.utils.quote(branch)}"
        f"&line={start_line}&lineEnd={end_line}"
        f"&lineStartColumn=1&lineEndColumn=25&type=2&lineStyle=plain&_a=files"
    )


def _build_ado_browse_line_url(file_path, branch, start_line, end_line):
    """Build an ADO file browse URL that highlights specific lines."""
    path_enc = requests.utils.quote(file_path)
    return (
        f"https://dev.azure.com/{ADO_ORG}/{ADO_PROJECT}/_git/{ADO_REPO}"
        f"?path={path_enc}&version=GB{requests.utils.quote(branch)}"
        f"&line={start_line}&lineEnd={end_line}"
        f"&lineStartColumn=1&lineEndColumn=25&lineStyle=plain&_a=contents"
    )


def fetch_recent_changes(iterations_log, branch):
    """Legacy fallback: fetch file-level commits (used by historical endpoint).

    Returns a list of dicts sorted by commit date (newest first):
      [{"file_path": "/src/Foo.cs", "file_name": "Foo.cs",
        "commits": [<commit_dict>, ...]}]
    """
    seen = set()
    file_paths = []
    for it in iterations_log:
        for tc in it.get("tool_calls", []):
            if tc.get("type") == "read_file" and tc.get("file_path") and not tc.get("error"):
                fp = tc["file_path"]
                if fp not in seen:
                    seen.add(fp)
                    file_paths.append(fp)

    if not file_paths:
        return []

    from concurrent.futures import ThreadPoolExecutor, as_completed
    results = {}

    def _fetch(fp):
        return fp, ado_get_file_commits(fp, branch, top=3)

    with ThreadPoolExecutor(max_workers=6) as pool:
        futures = {pool.submit(_fetch, fp): fp for fp in file_paths}
        for future in as_completed(futures):
            try:
                fp, commits = future.result()
                if commits:
                    results[fp] = commits
            except Exception:
                pass

    changes = []
    for fp, commits in results.items():
        file_name = fp.rsplit("/", 1)[-1] if "/" in fp else fp
        changes.append({
            "file_path": fp,
            "file_name": file_name,
            "commits":   commits,
        })

    changes.sort(
        key=lambda c: c["commits"][0]["date"] if c["commits"] else "",
        reverse=True,
    )
    return changes


def get_client():
    if not ANTHROPIC_API_KEY:
        raise ValueError("ANTHROPIC_API_KEY is not set.")
    return anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)




def _clean_answer(text: str) -> str:
    """Sanitise the user-facing answer by stripping internal directives and
    raw Python data structures that should never appear in output.

    Handles:
      - MISSING_FILE: directives (internal fetch signals)
      - Raw Python dict / list repr leaking from iterations_log formatting
        e.g.  {'type': 'read_file', 'file_path': '...', 'content_length': ...}
    """
    import re

    # ── Strip MISSING_FILE: lines ─────────────────────────────────────────────
    missing_files = [
        line.replace("MISSING_FILE:", "").strip()
        for line in text.splitlines()
        if "MISSING_FILE:" in line
    ]
    cleaned = re.sub(r"^MISSING_FILE:.*$\n?", "", text, flags=re.MULTILINE).strip()
    cleaned = re.sub(
        r"(?m)^(#+\s*Missing Files.*|What You Actually Need.*)\n+$", "", cleaned
    ).strip()

    # ── Strip raw Python dict/list tool-call representations ─────────────────
    # Matches lines like: {'type': 'read_file', 'file_path': '...', ...}
    cleaned = re.sub(
        r"^\{['\"]type['\"]\s*:.*\}\s*$\n?",
        "",
        cleaned,
        flags=re.MULTILINE,
    ).strip()
    # Also strip standalone lines that look like dict key-value pairs leaked
    # from str(iterations_log) formatting
    cleaned = re.sub(
        r"^['\"](?:type|file_path|content_length|query|results|preview)['\"].*$\n?",
        "",
        cleaned,
        flags=re.MULTILINE,
    ).strip()

    # ── Normalize deep headings (####+ → ###) for better cross-app rendering ──
    # Many external apps (Jira, Teams, embedded webviews) don't render h4+
    # headings, causing raw #### to appear as literal text.
    cleaned = re.sub(r"^#{4,}\s+", "### ", cleaned, flags=re.MULTILINE)

    # ── Add footnote if files were flagged as missing ─────────────────────────
    if missing_files:
        names = ", ".join(f"`{p.rsplit('/', 1)[-1]}`" for p in missing_files[:5])
        suffix = "…" if len(missing_files) > 5 else ""
        cleaned += (
            f"\n\n---\n> **Note:** A more complete answer may require additional "
            f"files that could not be located in the searched branch: {names}{suffix}"
        )
    return cleaned


def assess_confidence(query: str, answer: str, files_read: int, loop_name: str, model: str | None = None) -> dict:
    """Quick confidence assessment for a loop answer.
    Returns {"level": "High"|"Medium"|"Low", "reason": "one sentence"}
    """
    import json as _json
    client = get_client()
    try:
        resp = client.messages.create(
            model=model or MODEL,
            max_tokens=120,
            system=(
                "You evaluate the confidence level of Cora PPM search answers. "
                "Respond ONLY with a valid JSON object, no extra text:\n"
                '{"level": "High", "reason": "One sentence."}\n'
                "Levels:\n"
                "  High   — multiple relevant files read, answer covers all key layers\n"
                "  Medium — some relevant files found, coverage may be incomplete\n"
                "  Low    — limited or no matching files; answer is partial or inferred"
            ),
            messages=[{
                "role": "user",
                "content": (
                    f"Search method: {loop_name}\n"
                    f"Files read: {files_read}\n"
                    f"Question: {query}\n\n"
                    f"Answer preview (first 400 chars):\n{answer[:400]}\n\n"
                    "Rate the confidence in this answer."
                ),
            }],
        )
        text = resp.content[0].text.strip()
        # Strip markdown fences if present
        if "```" in text:
            text = text.split("```")[1]
            if text.startswith("json"):
                text = text[4:]
            text = text.strip()
        parsed = _json.loads(text)
        # Normalise level capitalisation
        parsed["level"] = parsed.get("level", "Medium").capitalize()
        return parsed
    except Exception:
        return {"level": "Medium", "reason": "Confidence could not be assessed automatically."}


def extract_file_paths_from_text(text: str) -> list:
    """Parse a synthesis response for any ADO-style file paths Claude mentioned.

    Looks for patterns like /code/inetpub/..., /src/..., or Helpers/Foo.cs
    so the agentic loop can fetch files Claude said it needed but didn't have.
    """
    import re
    # Match Unix-style paths that end with a known code extension
    pattern = re.compile(
        r'(?:^|[\s`\'"])(/[\w./ -]+\.(?:cs|vb|aspx|ascx|js|ts|cshtml|vbhtml|sql|config|xml|json))',
        re.MULTILINE | re.IGNORECASE,
    )
    paths = []
    seen  = set()
    for m in pattern.finditer(text):
        p = m.group(1).strip()
        if p not in seen:
            seen.add(p)
            paths.append(p)
    return paths


# ---------------------------------------------------------------------------
# 1. AGENTIC LOOP
#    Claude iteratively decides WHAT to search for and WHICH files to read
#    using two ADO-specific tools, until it has enough context to answer.
# ---------------------------------------------------------------------------

def run_agentic_loop(query, branch, job_id: str | None = None, model: str | None = None):
    _model = model or MODEL
    client = get_client()

    tools = [
        {
            "name": "ado_code_search",
            "description": (
                "Search the Cora PPM Azure DevOps repository for relevant code, "
                "files, or implementations. Returns matching file paths AND content "
                f"snippets (the actual matching lines of code) from the {ADO_REPO} repository. "
                "The snippets show the exact code at the match location — use them to "
                "quickly inspect code from large files without reading the full file. "
                "Call this multiple times with different short technical search terms "
                "(class names, method names, feature names, 1-3 words) to find all relevant files. "
                "NEVER use the full user question as the search term — it returns nothing."
            ),
            "input_schema": {
                "type": "object",
                "properties": {
                    "search_query": {
                        "type": "string",
                        "description": (
                            "Search terms to find relevant code or documentation "
                            "in the Cora PPM ADO repository."
                        ),
                    }
                },
                "required": ["search_query"],
            },
        },
        {
            "name": "ado_read_file",
            "description": (
                "Read the content of a specific file from the Cora PPM ADO repository. "
                "Use file paths returned by ado_code_search. "
                "Returns the FULL file content — no truncation. "
                "You will always receive the complete file in a single call."
            ),
            "input_schema": {
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": (
                            "Full path of the file in the repository, "
                            "e.g. /src/Features/Forecasting/EACSubmit.cs"
                        ),
                    },
                    "char_offset": {
                        "type": "integer",
                        "description": (
                            "Character offset to start reading from. "
                            "Omit or use 0 to read from the beginning."
                        ),
                    },
                },
                "required": ["file_path"],
            },
        },
    ]

    system_prompt = (
        f"You are a senior Cora PPM code analyst. Your job is to give COMPLETE, thorough answers "
        f"about the Cora PPM codebase — matching the quality of a developer who has read every file.\n\n"
        f"REPOSITORY: org={ADO_ORG}, project={ADO_PROJECT}, repo={ADO_REPO}, branch={branch}\n\n"
        "═══ SEARCHING STRATEGY ═══\n"
        "ADO Code Search does NOT understand natural language — long phrases return ZERO results.\n"
        "Use SHORT (1-3 word) technical terms. Search in LAYERS:\n"
        "  Layer 1 — The UI/page file: e.g. 'TimesheetGrid', 'EACSubmit', 'ForecastApproval'\n"
        "  Layer 2 — The short keyword: e.g. 'Timesheet', 'EAC', 'Forecast'\n"
        "  Layer 3 — The service/repo: e.g. 'ITimesheetService', 'TimesheetRepository'\n"
        "  Layer 4 — The controller/codebehind: e.g. 'TimesheetController', 'Timesheet.aspx'\n"
        "  Layer 5 — Alternate casing or legacy name: e.g. 'TimeSheet' if 'Timesheet' returned nothing\n"
        "Run AT LEAST 4 different search terms before concluding you have enough files.\n"
        "If a search returns 0 results, immediately try a shorter or differently-cased variant.\n\n"
        "═══ READING STRATEGY ═══\n"
        "Read files in this priority order and DO NOT stop after 2-3 files:\n"
        "  1. The .aspx / .ascx page or control file that renders the UI\n"
        "  2. The .aspx.vb or .aspx.cs code-behind that handles events\n"
        "  3. The service or business logic class (IXxxService, XxxService)\n"
        "  4. The repository or data-access class (XxxRepository, XxxDA)\n"
        "  5. Any helpers, utilities, or shared controls it calls\n"
        "  6. Any SQL stored procedures or queries referenced\n"
        "Read ALL layers — an answer that only covers the UI without the business logic is incomplete.\n\n"
        "═══ FILE READS — NO TRUNCATION ═══\n"
        "ado_read_file returns the FULL file content in one call — there is no truncation.\n"
        "You will always receive the complete file, so every method and class definition\n"
        "will be visible. Do NOT re-read the same file unless you need a different offset.\n\n"
        "═══ ANSWER REQUIREMENTS ═══\n"
        "Your answer MUST use EXACTLY this three-section structure with these EXACT headings:\n\n"
        "## Overview\n"
        "Write a complete, professional, authoritative plain-English explanation for the end user. "
        "Use plain language — no file paths, no code snippets, no method names. "
        "Explain what the feature/function does, why it exists, and how it works from the "
        "user's perspective. A non-technical stakeholder should fully understand this section.\n\n"
        "---\n\n"
        "## User Guide\n"
        "Provide a guided, step-by-step response that helps the user interact with or configure "
        "the feature. Include:\n"
        "  - Where to find the feature in the Cora PPM UI (navigation path)\n"
        "  - Step-by-step instructions for common tasks related to this feature\n"
        "  - Any important settings, options, or prerequisites the user should know about\n"
        "  - Tips, best practices, or common pitfalls to avoid\n"
        "Write as if you are guiding a user through the feature. Keep language clear and actionable.\n\n"
        "---\n\n"
        "## Technical Reference\n"
        "*For dev/product team review*\n\n"
        "Full technical detail for developers to verify:\n"
        "  - Complete execution flow from the ADO source (every method, file, layer in order).\n"
        "    After EACH step, add an inline annotation showing which Overview claim it explains:\n"
        "    e.g. `TimesheetService.SaveTimesheet()` in `TimesheetService.vb` (line 87) → *Supports: \"data is saved via a service layer\"*\n"
        "  - Key code snippets in fenced code blocks. Each snippet MUST be preceded by:\n"
        "    `File: path/to/File.ext, lines N–M` and annotated with which Overview claim it supports.\n"
        "  - File inventory with each file's role (UI / CodeBehind / Service / Repository / Helper / SQL) "
        "and which Overview statement(s) it backs.\n"
        "  - ⚠️ CONDITIONAL LOGIC & FEATURE TOGGLES: Explicitly call out ANY conditional branches,\n"
        "    feature flags, config-driven paths, or permission checks that affect the behaviour\n"
        "    described in the Overview. Format each as:\n"
        "    > ⚠️ **Conditional:** `If SomeFlag = True` in `File.ext` (line N) — affects: \"[Overview claim]\"\n"
        "  - Any gaps or missing layers identified during the search.\n\n"
        "═══ DISAMBIGUATION — SAME NAME, DIFFERENT UI ═══\n"
        "Cora PPM frequently reuses the same button label, feature name, or function name across "
        "DIFFERENT pages/UIs (e.g. 'Load From Schedule' exists on both the Resource Allocation page "
        "AND the Project Forecasting page, doing completely different things). When you encounter this:\n"
        "  1. SEARCH PHASE: If you discover that a feature name maps to multiple .aspx pages or "
        "     user controls, you MUST search and read EACH page's implementation separately.\n"
        "  2. ANSWER PHASE: NEVER merge or conflate implementations from different pages into one "
        "     explanation. Describe each page's version separately with clear headings:\n"
        "     e.g. '### On the Resource Allocation page' vs '### On the Project Forecasting page'\n"
        "  3. If the user's question specifies a particular page/UI, focus on that one but note "
        "     that the same feature name also exists elsewhere.\n"
        "  4. If the user's question does NOT specify which page, cover ALL pages where the feature "
        "     exists and clearly separate the explanations.\n\n"
        "NEVER use general knowledge. Only answer from files you have read. "
        "If you cannot find a file you need, say which file and why it matters, then continue "
        "with what you have rather than stopping early.\n\n"
        "═══ WHEN TO STOP SEARCHING ═══\n"
        "Do NOT stop searching until you have covered ALL of these layers:\n"
        "  ✓ The .aspx/.ascx page that renders the UI\n"
        "  ✓ The .aspx.vb/.aspx.cs code-behind that handles events\n"
        "  ✓ The service or business-logic class (IXxxService / XxxService)\n"
        "  ✓ The repository or data-access class (XxxRepository / XxxDA / XxxManager)\n"
        "  ✓ Any SQL stored procedures or queries called from the repository\n"
        "  ✓ Key helper / utility classes that the above files delegate to\n"
        "  ✓ Any base classes or shared controls that define important inherited behaviour\n\n"
        "You are NOT done until every layer above is checked. If a layer's file could not be "
        "found (e.g. no stored proc exists), note that explicitly and move on — do not stop early.\n\n"
        "AVOID reading: pure CSS/JS styling files, completely unrelated features, or files you "
        "have already read in full. Skip CSS/layout-only .ascx controls unless the question is "
        "specifically about layout.\n\n"
        "TARGET DEPTH: For most questions a thorough answer requires 12–20 files. If you have "
        "fewer than 10 files read and the question touches business logic, keep searching."
    )

    # Inject known-file hints from the codebase index (Approach 1+2)
    try:
        index_hits = lookup_codebase_index(query, branch, limit=15)
        index_hints = format_index_hints(index_hits)
        if index_hints:
            system_prompt += index_hints
            if job:
                job["event_queue"].put({
                    "type": "progress",
                    "text": f"Found {len(index_hits)} indexed file(s) from past searches — using as starting points",
                })
    except Exception as e:
        print(f"[Index] Warning: could not look up index: {e}")

    messages = [{"role": "user", "content": (
        f"Research and answer this question about the Cora PPM codebase.\n\n"
        f"Question: {query}\n\n"
        f"Remember: extract the KEY TECHNICAL TERMS first, then search with those short terms."
    )}]
    iterations_log      = []
    iteration           = 0
    final_answer        = ""
    total_input_tokens  = 0
    total_output_tokens = 0
    # Track full file contents separately so fallback synthesis has real context,
    # not just metadata dicts from the iterations log.
    accumulated_files = {}   # file_path -> full content string

    # Time-based control: hard stop and user-confirm thresholds
    loop_start_time     = time.time()
    next_pause_at       = loop_start_time + SEARCH_PAUSE_FIRST_SEC
    job                 = _job_get(job_id) if job_id else None

    while True:
        iteration += 1

        # ── Iteration hard cap ────────────────────────────────────────────────
        if iteration > 25:
            messages.append({"role": "user", "content": (
                "SYSTEM NOTE: You have reached the 25-iteration limit. "
                "Write your complete answer NOW using every file you have already read. "
                "Do not make any more tool calls."
            )})
            break

        # ── Time-based hard stop ──────────────────────────────────────────────
        elapsed = time.time() - loop_start_time
        if elapsed >= SEARCH_HARD_LIMIT_SEC:
            messages.append({"role": "user", "content": (
                "SYSTEM NOTE: The search has reached its 11-minute hard limit. "
                "Write your complete answer NOW using every file you have already read. "
                "Do not make any more tool calls."
            )})
            break

        # ── Pause threshold: ask user if they want to continue ────────────────
        if elapsed >= next_pause_at - loop_start_time and job:
            # Signal the frontend via SSE
            minutes_so_far = int(elapsed // 60)
            files_so_far   = len(accumulated_files)
            job["event_queue"].put({
                "type":    "pause_prompt",
                "elapsed": int(elapsed),
                "files":   files_so_far,
                "message": (
                    f"Search is still running ({minutes_so_far} min, {files_so_far} files read). "
                    "Continue searching for more files?"
                ),
            })
            # Wait up to 90 s for either a continue or stop signal from the frontend
            continue_evt = job["continue_evt"]
            stop_evt     = job["stop_evt"]
            continue_evt.clear()
            stop_evt.clear()
            deadline = time.time() + 90
            while time.time() < deadline:
                if stop_evt.is_set():
                    # User chose to stop — synthesise with what we have
                    messages.append({"role": "user", "content": (
                        "SYSTEM NOTE: The user has requested the search stop now. "
                        "Write your complete answer using every file you have already read. "
                        "Do not make any more tool calls."
                    )})
                    break
                if continue_evt.is_set():
                    next_pause_at = time.time() + SEARCH_PAUSE_REPEAT_SEC
                    break
                time.sleep(0.5)
            else:
                # Timeout waiting for user response — default to continue
                next_pause_at = time.time() + SEARCH_PAUSE_REPEAT_SEC
            if stop_evt.is_set():
                break

        response = client.messages.create(
            model=_model,
            max_tokens=5000,
            system=system_prompt,
            tools=tools,
            messages=messages,
        )
        total_input_tokens  += getattr(getattr(response, "usage", None), "input_tokens",  0)
        total_output_tokens += getattr(getattr(response, "usage", None), "output_tokens", 0)

        tool_uses = [b for b in response.content if b.type == "tool_use"]

        if not tool_uses:
            final_answer = _clean_answer(" ".join(
                b.text for b in response.content if hasattr(b, "text")
            ).strip())
            iterations_log.append({
                "iteration":  iteration,
                "action":     "Claude concluded — generating answer from gathered ADO context.",
                "tool_calls": [],
            })
            break

        tool_calls_this_turn = []
        tool_results         = []

        for tu in tool_uses:
            if tu.name == "ado_code_search":
                sq = tu.input.get("search_query", query)
                try:
                    results = ado_code_search(sq, branch)
                    result_content = json.dumps(results)
                    tool_calls_this_turn.append({
                        "type":    "search",
                        "query":   sq,
                        "results": results,
                    })
                except Exception as e:
                    result_content = json.dumps({"error": str(e)})
                    tool_calls_this_turn.append({"type": "search", "query": sq, "error": str(e)})

            elif tu.name == "ado_read_file":
                fp          = tu.input.get("file_path", "")
                char_offset = int(tu.input.get("char_offset") or 0)
                try:
                    content = ado_read_file(fp, branch, char_offset=char_offset)
                    result_content = content
                    # Continuation reads append to the existing entry so the
                    # full reconstructed file is available for synthesis.
                    if char_offset > 0 and fp in accumulated_files:
                        accumulated_files[fp] += f"\n\n[...continuing from offset {char_offset:,}...]\n\n" + content
                    else:
                        accumulated_files[fp] = content
                    tool_calls_this_turn.append({
                        "type":           "read_file",
                        "file_path":      fp,
                        "char_offset":    char_offset,
                        "content_length": len(content),
                        "preview":        content[:200] + "..." if len(content) > 200 else content,
                    })
                except Exception as e:
                    result_content = "Error reading file: " + str(e)
                    tool_calls_this_turn.append({"type": "read_file", "file_path": fp, "error": str(e)})

            else:
                result_content = "Unknown tool."

            tool_results.append({
                "type":        "tool_result",
                "tool_use_id": tu.id,
                "content":     result_content,
            })

        iterations_log.append({
            "iteration":  iteration,
            "action":     "Made " + str(len(tool_uses)) + " ADO tool call(s).",
            "tool_calls": tool_calls_this_turn,
        })

        # Emit a progress event so the frontend can show live step feedback
        if job:
            _prog_parts = []
            for _tc in tool_calls_this_turn:
                if _tc.get("type") == "search":
                    _prog_parts.append(f"🔎 Searching '{_tc.get('query', '')}'")
                elif _tc.get("type") == "read_file":
                    _fname = _tc.get("file_path", "").rsplit("/", 1)[-1]
                    _off   = _tc.get("char_offset", 0)
                    _prog_parts.append(
                        f"📖 Reading {_fname}" + (f" (+{_off:,})" if _off else "")
                    )
            _searches_so_far = sum(
                1 for _it in iterations_log
                for _tc2 in _it.get("tool_calls", [])
                if _tc2.get("type") == "search"
            ) + sum(1 for _tc2 in tool_calls_this_turn if _tc2.get("type") == "search")
            job["event_queue"].put({
                "type":      "progress",
                "iteration": iteration,
                "files":     len(accumulated_files),
                "searches":  _searches_so_far,
                "text":      " · ".join(_prog_parts) if _prog_parts else f"Iteration {iteration}",
            })

        messages.append({"role": "assistant", "content": response.content})
        messages.append({"role": "user",      "content": tool_results})

        if response.stop_reason == "end_turn":
            text = " ".join(b.text for b in response.content if hasattr(b, "text")).strip()
            if text:
                final_answer = _clean_answer(text)
            break

    # ── Fallback synthesis if iteration cap hit without final text ────────────
    # Build context from ACTUAL file contents (not metadata dicts) so Claude
    # can write a real answer rather than regurgitating internal tool call data.
    if not final_answer:
        if accumulated_files:
            context = "\n\n".join(
                f"### File: {path}\n{content}"
                for path, content in accumulated_files.items()
            )
        else:
            context = (
                "No files were successfully read from the repository. "
                "The searches ran but returned no accessible content."
            )
        fallback = client.messages.create(
            model=_model,
            max_tokens=10000,
            system=(
                "You are a senior Cora PPM code analyst. Answer ONLY from the ADO repository "
                "context provided — never from general knowledge.\n\n"
                "Your answer MUST use EXACTLY this three-section structure:\n\n"
                "## Overview\n"
                "Write a complete, professional, authoritative plain-English explanation for the end user. "
                "Use plain language — no file paths, no code snippets, no method names. "
                "Explain what the feature/function does, why it exists, and how it works from the "
                "user's perspective. A non-technical stakeholder should fully understand this section.\n\n"
                "---\n\n"
                "## User Guide\n"
                "Provide guided, step-by-step instructions for using or configuring the feature. "
                "Include navigation paths, prerequisites, tips, and common pitfalls.\n\n"
                "---\n\n"
                "## Technical Reference\n"
                "*For dev/product team review*\n\n"
                "Full technical detail for developers to verify:\n"
                "  - Complete execution flow (every method, file, layer in order) with inline annotations.\n"
                "  - Key code snippets in fenced code blocks with file paths.\n"
                "  - File inventory with each file's role and which Overview statement(s) it backs.\n"
                "  - Any conditional logic, feature flags, or permission checks.\n"
                "  - Any gaps or missing layers identified.\n\n"
                "DISAMBIGUATION: If the provided context contains implementations of the same feature "
                "name (button, function) from DIFFERENT pages or UIs, you MUST describe each separately "
                "with clear headings (e.g. '### On the Resource Allocation page' vs '### On the Project "
                "Forecasting page'). NEVER merge them into one explanation.\n\n"
                "If context is insufficient to cover all layers, state clearly which layer is missing "
                "and what additional files would be needed."
            ),
            messages=[{
                "role":    "user",
                "content": (
                    "Based ONLY on the following ADO repository context, answer in full detail:\n\n"
                    "Question: " + query + "\n\n"
                    "ADO context gathered:\n" + context
                ),
            }],
        )
        total_input_tokens  += getattr(getattr(fallback, "usage", None), "input_tokens",  0)
        total_output_tokens += getattr(getattr(fallback, "usage", None), "output_tokens", 0)
        final_answer = _clean_answer(fallback.content[0].text)

    total_searches   = sum(1 for it in iterations_log for tc in it.get("tool_calls", []) if tc.get("type") == "search")
    total_files_read = sum(1 for it in iterations_log for tc in it.get("tool_calls", []) if tc.get("type") == "read_file")

    confidence = assess_confidence(query, final_answer, total_files_read, "Agentic Loop", model=_model)

    # Post-processing: fetch blame-level recent commits for code sections Claude highlighted
    if job:
        job["event_queue"].put({"type": "progress", "text": "Fetching blame-level commit history for identified code sections…"})
    recent_changes = fetch_blame_changes(final_answer, accumulated_files, branch)

    # (Codebase index is now maintained via bulk import; hit_count bumped at lookup time)

    result = {
        "answer":           final_answer,
        "confidence":       confidence,
        "iterations":       iterations_log,
        "total_iterations": iteration,
        "total_searches":   total_searches,
        "total_files_read": total_files_read,
        "branch":           branch,
        "input_tokens":     total_input_tokens,
        "output_tokens":    total_output_tokens,
        "elapsed_seconds":  round(time.time() - loop_start_time, 1),
        "recent_changes":   recent_changes,
    }

    # Notify SSE stream that the job is done
    if job:
        job["event_queue"].put({"type": "done", "result": result})
    return result


# ---------------------------------------------------------------------------
# Flask routes
# ---------------------------------------------------------------------------

@app.route("/")
@login_required
def index():
    user = session.get("user", {}) if not AUTH_DISABLED else {}
    return render_template("index.html",
                           user=user,
                           is_admin=_is_admin(),
                           auth_disabled=AUTH_DISABLED,
                           cost_per_m_input=COST_PER_M_INPUT,
                           cost_per_m_output=COST_PER_M_OUTPUT,
                           model_pricing=MODEL_PRICING)


@app.route("/api-docs")
@login_required
def api_docs():
    """Interactive Swagger-style API documentation page."""
    user = session.get("user", {}) if not AUTH_DISABLED else {}
    return render_template("api_docs.html",
                           user=user,
                           is_admin=_is_admin(),
                           auth_disabled=AUTH_DISABLED)


def _extract_screenshot_context(image_b64: str, mime_type: str = "image/png", model: str | None = None) -> str:
    """Ask Claude to extract UI context from a base64-encoded screenshot.

    Returns the extracted plain-text context string, or raises on failure.
    """
    client = get_client()
    resp = client.messages.create(
        model=model or MODEL,
        max_tokens=600,
        system=(
            "You are a Cora PPM UI analyst. The user has uploaded a screenshot of the Cora PPM application. "
            "Extract ALL visible context that would help a developer understand what page/feature is shown. "
            "Respond ONLY with a concise plain-text summary (no markdown) covering:\n"
            "• Page name / module visible in the screenshot\n"
            "• Key field labels, button names, or column headers visible\n"
            "• Any error messages, warning banners, or validation messages\n"
            "• Any IDs, statuses, or data values that appear to be relevant\n"
            "• Any UI elements or controls that look directly related to the user's question\n"
            "Keep the response under 250 words. Do not include generic observations."
        ),
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {
                        "type":       "base64",
                        "media_type": mime_type,
                        "data":       image_b64,
                    },
                },
                {
                    "type": "text",
                    "text": "Please extract the UI context from this Cora PPM screenshot.",
                },
            ],
        }],
    )
    return resp.content[0].text.strip()


@app.route("/api/extract-screenshot", methods=["POST"])
def api_extract_screenshot():
    """
    Receive a base64-encoded screenshot image, ask Claude to extract UI context
    (page name, visible fields, error messages, data shown), and return the
    context as a plain-text string to be appended to the user's query.
    """
    data      = request.get_json() or {}
    image_b64 = data.get("image_b64", "").strip()
    mime_type = data.get("mime_type", "image/png").strip()
    _model    = _resolve_model(data)
    if not image_b64:
        return jsonify({"error": "No image data provided."}), 400
    try:
        context = _extract_screenshot_context(image_b64, mime_type, model=_model)
        return jsonify({"success": True, "context": context})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/branches")
def api_branches():
    try:
        branches = ado_get_branches()
        return jsonify({"success": True, "branches": branches})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/test")
def api_test():
    """Diagnostic endpoint — tests ADO auth, branch listing, and a sample search."""
    results = {}

    # 1. Check env vars
    results["anthropic_key_set"] = bool(os.environ.get("ANTHROPIC_API_KEY"))
    results["ado_pat_set"]       = bool(os.environ.get("AZURE_DEVOPS_PAT"))

    # 2. Test branch listing
    try:
        branches = ado_get_branches()
        results["branches_ok"]    = True
        results["branch_count"]   = len(branches)
        results["branches_sample"] = branches[:3]
    except Exception as e:
        results["branches_ok"]    = False
        results["branches_error"] = str(e)

    # 3. Test a minimal code search against the first branch found
    test_branch = branches[0] if results.get("branches_ok") and branches else ""
    if test_branch:
        try:
            hits = ado_code_search("Forecasting", test_branch, top=2)
            results["search_ok"]      = True
            results["search_branch"]  = test_branch
            results["search_hits"]    = hits
        except Exception as e:
            results["search_ok"]      = False
            results["search_branch"]  = test_branch
            results["search_error"]   = str(e)
    else:
        results["search_ok"]    = False
        results["search_error"] = "No branch available to test against"

    return jsonify(results)


@app.route("/api/agentic/stream", methods=["POST"])
def api_agentic_stream():
    """
    SSE endpoint for the Agentic Loop.
    - Starts the loop in a background thread.
    - Streams progress events (pause_prompt, done, error) as SSE.
    - Returns a job_id in the first event so the frontend can call
      /api/agentic/continue or /api/agentic/stop.
    """
    data   = request.get_json() or {}
    query  = data.get("query",  "").strip()
    branch = data.get("branch", "").strip()
    _model = _resolve_model(data)
    if not query:
        return jsonify({"error": "No query provided."}), 400
    if not branch:
        return jsonify({"error": "Please select a branch first."}), 400

    # Capture screenshot data so it can be persisted with the search event
    _ss_b64  = data.get("screenshot_b64", "").strip()
    _ss_mime = data.get("screenshot_mime", "").strip()
    _ss_ctx  = data.get("screenshot_context", "").strip()

    job_id = str(uuid.uuid4())
    _job_create(job_id)
    job    = _job_get(job_id)
    t0     = time.time()
    _user_email = _current_user_email()

    # Persist job to api_jobs so GET /api/query/<job_id> works for all searches
    _persist_api_job(job_id, query, branch, _user_email)

    def _run():
        try:
            result = run_agentic_loop(query, branch, job_id=job_id, model=_model)
            event_id = track_search_event(
                "agentic", query, branch, "success",
                time.time() - t0,
                result.get("total_files_read", 0),
                result.get("total_iterations",  0),
                result.get("total_searches",    0),
                result.get("confidence", {}).get("level", ""),
                input_tokens=result.get("input_tokens",  0),
                output_tokens=result.get("output_tokens", 0),
                user_email=_user_email,
                answer=result.get("answer", ""),
                job_id=job_id,
                screenshot_context=_ss_ctx,
                screenshot_b64=_ss_b64,
                screenshot_mime=_ss_mime,
                model=_model,
            )
            if event_id:
                result["_event_id"] = event_id
            _complete_api_job(job_id, result)
        except Exception as e:
            track_search_event("agentic", query, branch, "error", time.time() - t0,
                               error_message=str(e), user_email=_user_email, job_id=job_id,
                               model=_model)
            _fail_api_job(job_id, str(e), time.time() - t0)
            job["event_queue"].put({"type": "error", "message": str(e)})
        finally:
            _job_remove(job_id)

    threading.Thread(target=_run, daemon=True).start()

    def _generate():
        # First event: send the job_id to the frontend
        yield f"data: {json.dumps({'type': 'started', 'job_id': job_id})}\n\n"
        while True:
            try:
                evt = job["event_queue"].get(timeout=30)
                yield f"data: {json.dumps(evt)}\n\n"
                if evt.get("type") in ("done", "error"):
                    break
            except queue.Empty:
                # Keep-alive ping
                yield "data: {\"type\":\"ping\"}\n\n"

    return Response(
        stream_with_context(_generate()),
        mimetype="text/event-stream",
        headers={
            "Cache-Control":        "no-cache",
            "X-Accel-Buffering":    "no",
            "Access-Control-Allow-Origin": "*",
        },
    )


@app.route("/api/agentic/start", methods=["POST"])
def api_agentic_start():
    """
    Polling-friendly replacement for /api/agentic/stream.
    Starts the agentic loop in a background thread and immediately returns
    a job_id.  The frontend then polls /api/agentic/poll/<job_id> every ~2 s.
    This avoids Railway's ~300 s HTTP proxy timeout that kills long SSE streams.
    """
    data   = request.get_json() or {}
    query  = data.get("query",  "").strip()
    branch = data.get("branch", "").strip()
    _model = _resolve_model(data)
    if not query:
        return jsonify({"error": "No query provided."}), 400
    if not branch:
        return jsonify({"error": "Please select a branch first."}), 400

    # Capture screenshot data so it can be persisted with the search event
    _ss_b64  = data.get("screenshot_b64", "").strip()
    _ss_mime = data.get("screenshot_mime", "").strip()
    _ss_ctx  = data.get("screenshot_context", "").strip()

    job_id = str(uuid.uuid4())
    _job_create(job_id)
    job = _job_get(job_id)
    t0  = time.time()
    _user_email = _current_user_email()

    # Persist job to api_jobs so GET /api/query/<job_id> works for all searches
    _persist_api_job(job_id, query, branch, _user_email)

    def _run():
        try:
            result = run_agentic_loop(query, branch, job_id=job_id, model=_model)
            event_id = track_search_event(
                "agentic", query, branch, "success",
                time.time() - t0,
                result.get("total_files_read", 0),
                result.get("total_iterations",  0),
                result.get("total_searches",    0),
                result.get("confidence", {}).get("level", ""),
                input_tokens=result.get("input_tokens",  0),
                output_tokens=result.get("output_tokens", 0),
                user_email=_user_email,
                answer=result.get("answer", ""),
                job_id=job_id,
                screenshot_context=_ss_ctx,
                screenshot_b64=_ss_b64,
                screenshot_mime=_ss_mime,
                model=_model,
            )
            if event_id:
                result["_event_id"] = event_id
            _complete_api_job(job_id, result)
        except Exception as e:
            track_search_event("agentic", query, branch, "error",
                               time.time() - t0, error_message=str(e),
                               user_email=_user_email, job_id=job_id,
                               model=_model)
            _fail_api_job(job_id, str(e), time.time() - t0)
            job["event_queue"].put({"type": "error", "message": str(e)})
        finally:
            # Keep job alive for 5 minutes after completion so the frontend
            # can drain the last events even if a poll arrives a bit late.
            threading.Timer(300, lambda: _job_remove(job_id)).start()

    threading.Thread(target=_run, daemon=True).start()
    return jsonify({"job_id": job_id})


@app.route("/api/agentic/poll/<job_id>", methods=["GET"])
def api_agentic_poll(job_id):
    """
    Drain all queued events for a running/completed agentic job.
    Returns {"events": [...]} — frontend calls this every ~2 s.
    Stops polling when it receives a "done" or "error" event.
    """
    job = _job_get(job_id)
    if not job:
        return jsonify({"events": [{"type": "error",
                                    "message": "Job not found or expired."}]})
    events = []
    try:
        while True:
            evt = job["event_queue"].get_nowait()
            events.append(evt)
            if evt.get("type") in ("done", "error"):
                break
    except queue.Empty:
        pass
    return jsonify({"events": events})


@app.route("/api/agentic/continue", methods=["POST"])
def api_agentic_continue():
    job_id = (request.get_json() or {}).get("job_id", "")
    job    = _job_get(job_id)
    if not job:
        return jsonify({"error": "Job not found."}), 404
    job["continue_evt"].set()
    return jsonify({"ok": True})


@app.route("/api/agentic/stop", methods=["POST"])
def api_agentic_stop():
    job_id = (request.get_json() or {}).get("job_id", "")
    job    = _job_get(job_id)
    if not job:
        return jsonify({"error": "Job not found."}), 404
    job["stop_evt"].set()
    return jsonify({"ok": True})


@app.route("/api/agentic", methods=["POST"])
def api_agentic():
    """Legacy synchronous endpoint — kept for backwards compatibility."""
    data   = request.get_json() or {}
    query  = data.get("query",  "").strip()
    branch = data.get("branch", "").strip()
    _model = _resolve_model(data)
    if not query:
        return jsonify({"error": "No query provided."}), 400
    if not branch:
        return jsonify({"error": "Please select a branch first."}), 400
    _ss_b64  = data.get("screenshot_b64", "").strip()
    _ss_mime = data.get("screenshot_mime", "").strip()
    _ss_ctx  = data.get("screenshot_context", "").strip()
    t0 = time.time()
    _user_email = _current_user_email()
    job_id = str(uuid.uuid4())
    _persist_api_job(job_id, query, branch, _user_email)
    try:
        result = run_agentic_loop(query, branch, model=_model)
        event_id = track_search_event(
            "agentic", query, branch, "success",
            time.time() - t0,
            result.get("total_files_read", 0),
            result.get("total_iterations",  0),
            result.get("total_searches",    0),
            result.get("confidence", {}).get("level", ""),
            input_tokens=result.get("input_tokens",  0),
            output_tokens=result.get("output_tokens", 0),
            user_email=_user_email,
            answer=result.get("answer", ""),
            job_id=job_id,
            screenshot_context=_ss_ctx,
            screenshot_b64=_ss_b64,
            screenshot_mime=_ss_mime,
            model=_model,
        )
        if event_id:
            result["_event_id"] = event_id
        _complete_api_job(job_id, result)
        return jsonify({"success": True, "result": result})
    except Exception as e:
        track_search_event("agentic", query, branch, "error", time.time() - t0,
                           error_message=str(e), user_email=_user_email, job_id=job_id,
                           model=_model)
        _fail_api_job(job_id, str(e), time.time() - t0)
        return jsonify({"error": str(e)}), 500










# ---------------------------------------------------------------------------
# Clarifying question
# ---------------------------------------------------------------------------

@app.route("/api/clarify", methods=["POST"])
def api_clarify():
    """
    Quickly check whether a query is specific enough, or needs one clarifying
    question before searching. Returns JSON:
      {"needs_clarification": false}
    or
      {"needs_clarification": true, "question": "..."}
    """
    data  = request.get_json() or {}
    query = data.get("query", "").strip()
    _model = _resolve_model(data)
    if not query:
        return jsonify({"needs_clarification": False})
    try:
        client = get_client()
        resp   = client.messages.create(
            model=_model,
            max_tokens=200,
            system=(
                "You decide whether a Cora PPM codebase question needs one clarifying question "
                "before searching, or is already specific enough.\n\n"
                "Rules:\n"
                "  - Only ask if the answer would meaningfully change WHAT files to search for.\n"
                "  - Never ask about things that don't affect the search (e.g. 'why do you want to know').\n"
                "  - If the question names a specific feature, button, page, or behaviour → it's specific enough.\n"
                "  - Respond ONLY with valid JSON, no extra text:\n"
                '    {"needs_clarification": false}\n'
                "  or:\n"
                '    {"needs_clarification": true, "question": "One short, specific follow-up question."}'
            ),
            messages=[{"role": "user", "content": f"Question: {query}"}],
        )
        text = resp.content[0].text.strip()
        if "```" in text:
            text = text.split("```")[1]
            if text.startswith("json"): text = text[4:]
            text = text.strip()
        result = json.loads(text)
        result["needs_clarification"] = bool(result.get("needs_clarification", False))
        return jsonify(result)
    except Exception:
        return jsonify({"needs_clarification": False})


# ---------------------------------------------------------------------------
# Browse past queries (available to all authenticated users)
# ---------------------------------------------------------------------------

@app.route("/api/browse-queries", methods=["GET"])
def api_browse_queries():
    """
    Return recent successful search events for any user to browse.
    Lightweight — no answer text, no screenshots.  Just enough to power
    a 'Past Queries' table the user can click through.

    Query params:
      page       (int, default 1)
      limit      (int, default 25, max 100)
      search     (str, optional — filters query text)
      branch         (str, optional — exact branch filter)
      confidence     (str, optional — exact confidence_level filter)
      has_screenshot (str, optional — "yes" or "no")
    """
    if not DATABASE_URL:
        return jsonify({"rows": [], "total": 0})
    page  = max(1, int(request.args.get("page", 1)))
    limit = min(int(request.args.get("limit", 25)), 100)
    search         = (request.args.get("search") or "").strip()
    branch         = (request.args.get("branch") or "").strip()
    confidence     = (request.args.get("confidence") or "").strip()
    has_screenshot = (request.args.get("has_screenshot") or "").strip().lower()
    offset = (page - 1) * limit

    conn = None
    try:
        conn = get_db_conn()
        cur  = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

        where = "WHERE status = 'success' AND answer_text IS NOT NULL AND answer_text != ''"
        params = []
        if search:
            where += " AND query ILIKE %s"
            params.append(f"%{search}%")
        if branch:
            where += " AND branch = %s"
            params.append(branch)
        if confidence:
            where += " AND LOWER(confidence_level) = LOWER(%s)"
            params.append(confidence)
        if has_screenshot == 'yes':
            where += " AND COALESCE(screenshot_b64, '') != ''"
        elif has_screenshot == 'no':
            where += " AND (screenshot_b64 IS NULL OR screenshot_b64 = '')"

        # Total count
        cur.execute(f"SELECT COUNT(*) AS cnt FROM search_events {where}", params)
        total = cur.fetchone()["cnt"]

        # Page of rows
        cur.execute(f"""
            SELECT id, query, branch, ran_at, confidence_level,
                   duration_sec, files_read, iterations, searches,
                   COALESCE(user_email, '') AS user_email,
                   COALESCE(model, '') AS model,
                   (COALESCE(screenshot_b64, '') != '') AS has_screenshot
            FROM   search_events {where}
            ORDER BY ran_at DESC
            LIMIT %s OFFSET %s
        """, params + [limit, offset])
        rows = [dict(r) for r in cur.fetchall()]
        for r in rows:
            if r.get("ran_at"):
                r["ran_at"] = r["ran_at"].isoformat()

        # Distinct filter options (unfiltered) for populating dropdowns
        cur.execute("""
            SELECT DISTINCT branch FROM search_events
            WHERE status='success' AND answer_text IS NOT NULL AND answer_text != ''
            AND branch IS NOT NULL AND branch != ''
            ORDER BY branch
        """)
        branches = [r["branch"] for r in cur.fetchall()]

        cur.execute("""
            SELECT DISTINCT confidence_level FROM search_events
            WHERE status='success' AND answer_text IS NOT NULL AND answer_text != ''
            AND confidence_level IS NOT NULL AND confidence_level != ''
            ORDER BY confidence_level
        """)
        confidences = [r["confidence_level"] for r in cur.fetchall()]

        return jsonify({"rows": rows, "total": total, "page": page, "limit": limit,
                        "branches": branches, "confidences": confidences})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


# ---------------------------------------------------------------------------
# Similar-query lookup  (cost-saver — intercepts duplicate searches)
# ---------------------------------------------------------------------------

@app.route("/api/similar-queries", methods=["POST"])
def api_similar_queries():
    """
    Check for past search events whose query is similar to the incoming one.
    Uses PostgreSQL pg_trgm trigram similarity.  Returns up to 5 matches with
    a short answer preview so the user can decide whether to reuse a past result.

    Request JSON:  {"query": "...", "branch": "..." (optional)}
    Response JSON: {"matches": [{ id, query, branch, ran_at, confidence_level,
                                   answer_preview, similarity }]}
    """
    data  = request.get_json() or {}
    query = data.get("query", "").strip()
    branch = data.get("branch", "").strip()
    # Strip screenshot context from the incoming query too — compare only user text
    sc_idx = query.find("[Screenshot context")
    if sc_idx > 0:
        query = query[:sc_idx].strip()
    if not query or not DATABASE_URL:
        return jsonify({"matches": []})

    conn = None
    try:
        conn = get_db_conn()
        cur  = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

        # Similarity threshold — only return reasonably close matches
        threshold = 0.25

        # Strip screenshot context from stored queries before comparing.
        # The query column may contain appended "[Screenshot context ...]" blobs
        # which dilute trigram similarity when comparing against the user's text.
        clean_q = """
            CASE WHEN query LIKE '%%[Screenshot context%%'
                 THEN TRIM(LEFT(query, POSITION('[Screenshot context' IN query) - 1))
                 ELSE query
            END
        """

        sql = f"""
            SELECT id, query, branch, ran_at, confidence_level,
                   LEFT(answer_text, 500) AS answer_preview,
                   similarity({clean_q}, %s) AS sim,
                   (COALESCE(screenshot_b64, '') != '') AS has_screenshot
            FROM   search_events
            WHERE  status = 'success'
              AND  answer_text IS NOT NULL AND answer_text != ''
              AND  similarity({clean_q}, %s) > %s
        """
        params = [query, query, threshold]

        # Optionally filter to same branch
        if branch:
            sql += " AND branch = %s"
            params.append(branch)

        sql += " ORDER BY sim DESC, ran_at DESC LIMIT 5"

        cur.execute(sql, params)
        rows = [dict(r) for r in cur.fetchall()]

        # Serialise datetimes
        for row in rows:
            if row.get("ran_at"):
                row["ran_at"] = row["ran_at"].isoformat()
            row["similarity"] = round(float(row.get("sim", 0)), 3)
            row.pop("sim", None)

        return jsonify({"matches": rows})
    except Exception as e:
        print(f"[similar-queries] Error: {e}")
        return jsonify({"matches": []})
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/search-events/<int:event_id>/answer", methods=["GET"])
def api_search_event_answer(event_id):
    """Return the full answer_text for a specific search event (for loading past results)."""
    if not DATABASE_URL:
        return jsonify({"error": "No database configured"}), 500
    conn = None
    try:
        conn = get_db_conn()
        cur  = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(
            """SELECT id, query, branch, ran_at, confidence_level, answer_text,
                      response_sections, duration_sec, files_read, iterations,
                      searches, input_tokens, output_tokens, total_tokens,
                      COALESCE(screenshot_context, '') AS screenshot_context,
                      (COALESCE(screenshot_b64, '') != '') AS has_screenshot
               FROM search_events WHERE id = %s AND status = 'success'""",
            (event_id,),
        )
        row = cur.fetchone()
        if not row:
            return jsonify({"error": "Event not found"}), 404
        row = dict(row)
        if row.get("ran_at"):
            row["ran_at"] = row["ran_at"].isoformat()
        return jsonify(row)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/search-events/<int:event_id>/recent-changes", methods=["GET"])
def api_search_event_recent_changes(event_id):
    """Fetch blame-level recent commits for code sections in a past search result.

    Parses code sections with line ranges from the stored answer_text and
    fetches commits for those specific sections. Falls back to file-level
    commits if no code sections are found.
    """
    if not DATABASE_URL:
        return jsonify({"error": "No database configured"}), 500
    conn = None
    try:
        conn = get_db_conn()
        cur  = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(
            "SELECT answer_text, branch FROM search_events WHERE id = %s AND status = 'success'",
            (event_id,),
        )
        row = cur.fetchone()
        if not row:
            return jsonify({"error": "Event not found"}), 404
        answer = row.get("answer_text") or ""
        branch = row.get("branch") or "main"
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)

    # Try blame-level extraction first (no accumulated_files for historical results)
    changes = fetch_blame_changes(answer, {}, branch)

    # Fallback to file-level if no code sections were parsed
    if not changes:
        import re
        path_pattern = re.compile(r'(?:^|[`\s(])(/[A-Za-z0-9_./-]+\.[a-zA-Z]{1,10})(?:[`\s),;:]|$)', re.MULTILINE)
        paths = list(dict.fromkeys(m for m in path_pattern.findall(answer) if len(m) > 5))
        if paths:
            fake_log = [{"tool_calls": [{"type": "read_file", "file_path": p} for p in paths]}]
            changes = fetch_recent_changes(fake_log, branch)

    return jsonify({"recent_changes": changes})


@app.route("/api/search-events/<int:event_id>/screenshot", methods=["GET"])
def api_search_event_screenshot(event_id):
    """Return the stored screenshot image for a search event as a data-URI JSON response."""
    if not DATABASE_URL:
        return jsonify({"error": "No database configured"}), 500
    conn = None
    try:
        conn = get_db_conn()
        cur  = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(
            """SELECT COALESCE(screenshot_b64, '') AS screenshot_b64,
                      COALESCE(screenshot_mime, 'image/png') AS screenshot_mime
               FROM search_events WHERE id = %s""",
            (event_id,),
        )
        row = cur.fetchone()
        if not row or not row["screenshot_b64"]:
            return jsonify({"error": "No screenshot stored for this event"}), 404
        return jsonify({
            "screenshot_b64": row["screenshot_b64"],
            "screenshot_mime": row["screenshot_mime"],
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Export API — generate Word (.docx) or PDF from search results
# ---------------------------------------------------------------------------

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import markdown as md_lib
import re as _re

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph, Spacer, Table as RLTable, TableStyle, HRFlowable
from reportlab.lib import colors as rl_colors


def _md_to_plain(text):
    """Strip markdown to plain text (rough but effective)."""
    text = _re.sub(r'```[\s\S]*?```', lambda m: m.group(0).strip('`').strip(), text)
    text = _re.sub(r'`([^`]+)`', r'\1', text)
    text = _re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = _re.sub(r'\*([^*]+)\*', r'\1', text)
    text = _re.sub(r'^#{1,6}\s+', '', text, flags=_re.MULTILINE)
    text = _re.sub(r'^\s*[-*]\s+', '  • ', text, flags=_re.MULTILINE)
    text = _re.sub(r'^\s*\d+\.\s+', '  ', text, flags=_re.MULTILINE)
    text = _re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text.strip()


def _build_docx(query, branch, metadata, sections):
    """Build a .docx buffer from search result sections."""
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1A, 0x25, 0x35)

    # Title
    title = doc.add_heading('Cora Product Investigator', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x25, 0x35)

    # Query
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(query)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Metadata row
    meta_parts = []
    if branch:
        meta_parts.append(f"Branch: {branch}")
    if metadata.get('confidence'):
        meta_parts.append(f"Confidence: {metadata['confidence']}")
    if metadata.get('elapsed_seconds'):
        meta_parts.append(f"Time: {metadata['elapsed_seconds']}s")
    if metadata.get('total_tokens'):
        meta_parts.append(f"Tokens: {metadata['total_tokens']:,}")
    if meta_parts:
        mp = doc.add_paragraph(' | '.join(meta_parts))
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in mp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph('')  # spacer

    # Sections
    section_config = [
        ('overview', 'Overview', RGBColor(0x25, 0x63, 0xEB)),
        ('user_guide', 'User Guide', RGBColor(0x16, 0xA3, 0x4A)),
        ('tech_ref', 'Technical Reference', RGBColor(0xD9, 0x77, 0x06)),
    ]

    for key, label, color in section_config:
        content = sections.get(key, '').strip()
        if not content:
            continue

        heading = doc.add_heading(label, level=1)
        for run in heading.runs:
            run.font.color.rgb = color

        # Add content as paragraphs
        plain = _md_to_plain(content)
        for line in plain.split('\n'):
            line = line.rstrip()
            if not line:
                doc.add_paragraph('')
            elif line.startswith('  • '):
                doc.add_paragraph(line[4:], style='List Bullet')
            else:
                doc.add_paragraph(line)

    # Footer
    fp = doc.add_paragraph('')
    fp.add_run('Generated by Cora Product Investigator').font.size = Pt(8)
    fp.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _build_pdf(query, branch, metadata, sections):
    """Build a PDF buffer from search result sections."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            topMargin=0.75*inch, bottomMargin=0.75*inch,
                            leftMargin=0.75*inch, rightMargin=0.75*inch)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='DocTitle', fontName='Helvetica-Bold', fontSize=20,
                              textColor=HexColor('#1A2535'), alignment=1, spaceAfter=6))
    styles.add(ParagraphStyle(name='DocQuery', fontName='Helvetica', fontSize=12,
                              textColor=HexColor('#64748B'), alignment=1, spaceAfter=4))
    styles.add(ParagraphStyle(name='DocMeta', fontName='Helvetica', fontSize=8,
                              textColor=HexColor('#94A3B8'), alignment=1, spaceAfter=16))
    styles.add(ParagraphStyle(name='SectionHead', fontName='Helvetica-Bold', fontSize=14,
                              spaceBefore=18, spaceAfter=8))
    styles.add(ParagraphStyle(name='BodyText2', fontName='Helvetica', fontSize=10,
                              textColor=HexColor('#1A2535'), leading=14, spaceAfter=4))
    styles.add(ParagraphStyle(name='BulletItem', fontName='Helvetica', fontSize=10,
                              textColor=HexColor('#1A2535'), leading=14,
                              leftIndent=18, bulletIndent=6, spaceAfter=3))
    styles.add(ParagraphStyle(name='Footer', fontName='Helvetica', fontSize=7,
                              textColor=HexColor('#94A3B8'), alignment=1, spaceBefore=20))

    elements = []
    elements.append(RLParagraph('Cora Product Investigator', styles['DocTitle']))
    elements.append(RLParagraph(query.replace('&', '&amp;').replace('<', '&lt;'), styles['DocQuery']))

    meta_parts = []
    if branch:
        meta_parts.append(f"Branch: {branch}")
    if metadata.get('confidence'):
        meta_parts.append(f"Confidence: {metadata['confidence']}")
    if metadata.get('elapsed_seconds'):
        meta_parts.append(f"Time: {metadata['elapsed_seconds']}s")
    if metadata.get('total_tokens'):
        meta_parts.append(f"Tokens: {metadata['total_tokens']:,}")
    if meta_parts:
        elements.append(RLParagraph(' | '.join(meta_parts), styles['DocMeta']))

    elements.append(HRFlowable(width="100%", thickness=1, color=HexColor('#E2E8F0'), spaceAfter=12))

    section_config = [
        ('overview', 'Overview', '#2563EB'),
        ('user_guide', 'User Guide', '#16A34A'),
        ('tech_ref', 'Technical Reference', '#D97706'),
    ]

    for key, label, color in section_config:
        content = sections.get(key, '').strip()
        if not content:
            continue

        head_style = ParagraphStyle(f'Head_{key}', parent=styles['SectionHead'],
                                    textColor=HexColor(color))
        elements.append(RLParagraph(label, head_style))

        plain = _md_to_plain(content)
        for line in plain.split('\n'):
            line = line.rstrip()
            if not line:
                elements.append(Spacer(1, 6))
            elif line.startswith('  • '):
                safe = line[4:].replace('&', '&amp;').replace('<', '&lt;')
                elements.append(RLParagraph(f'• {safe}', styles['BulletItem']))
            else:
                safe = line.replace('&', '&amp;').replace('<', '&lt;')
                elements.append(RLParagraph(safe, styles['BodyText2']))

    elements.append(RLParagraph('Generated by Cora Product Investigator', styles['Footer']))

    doc.build(elements)
    buf.seek(0)
    return buf


@app.route("/api/export", methods=["POST"])
@login_required
def api_export():
    """Generate a Word (.docx) or PDF export of the search results.

    JSON body:
        format: "docx" or "pdf"
        query: the search query
        branch: branch searched
        metadata: { confidence, elapsed_seconds, total_tokens, ... }
        sections: { overview: "...", user_guide: "...", tech_ref: "..." }
    """
    data = request.get_json(force=True) or {}
    fmt       = data.get("format", "docx").lower()
    query     = data.get("query", "Search Results")
    branch    = data.get("branch", "")
    metadata  = data.get("metadata", {})
    sections  = data.get("sections", {})

    if fmt not in ("docx", "pdf"):
        return jsonify({"error": "format must be 'docx' or 'pdf'"}), 400

    try:
        if fmt == "docx":
            buf = _build_docx(query, branch, metadata, sections)
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ext = "docx"
        else:
            buf = _build_pdf(query, branch, metadata, sections)
            mimetype = "application/pdf"
            ext = "pdf"

        slug = _re.sub(r'[^a-z0-9]+', '-', query[:60].lower()).strip('-') or 'search-results'
        filename = f"cora-investigator-{slug}.{ext}"

        return Response(
            buf.getvalue(),
            mimetype=mimetype,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ---------------------------------------------------------------------------
# Flag API — users flag unhelpful / missing responses
# ---------------------------------------------------------------------------

@app.route("/api/flag", methods=["POST"])
@login_required
def api_flag():
    """Submit a unified rating for a search result (one per user per search_event_id).

    JSON body:
        search_event_id (int) — the search event being rated
        section_ratings (dict) — e.g. {"overview":"accurate","user_guide":"accurate","tech_ref":"inaccurate"}
        explanation (str) — required if any section is marked inaccurate
        request_article (bool) — whether to request a Confluence article
        query, loop_type, answer, branch, confidence_level, duration_sec,
        input_tokens, output_tokens, total_tokens, files_read, iterations, searches
    Also supports legacy per-container format (flag_type + container).
    """
    data = request.get_json() or {}
    query       = data.get("query", "").strip()
    loop_type   = data.get("loop_type", "agentic").strip()
    answer      = data.get("answer", "").strip()
    explanation = data.get("explanation", "").strip()
    request_article = 1 if data.get("request_article") else 0
    search_event_id = data.get("search_event_id")

    # New unified format: section_ratings dict
    section_ratings = data.get("section_ratings", {})
    if section_ratings:
        has_inaccurate = any(v == "inaccurate" for v in section_ratings.values())
        if has_inaccurate and not explanation:
            return jsonify({"error": "Please explain what was inaccurate or incomplete."}), 400
        # Derive overall flag_type: inaccurate if ANY section is inaccurate
        flag_type = "inaccurate" if has_inaccurate else "accurate"
        container = ",".join(sorted(section_ratings.keys()))
    else:
        # Legacy per-container format
        flag_type   = data.get("flag_type", "inaccurate").strip()
        container   = data.get("container", "").strip()

    if not query and not search_event_id:
        return jsonify({"error": "query or search_event_id is required"}), 400
    if flag_type not in ("accurate", "inaccurate"):
        return jsonify({"error": "flag_type must be 'accurate' or 'inaccurate'"}), 400
    if flag_type == "inaccurate" and not explanation:
        return jsonify({"error": "explanation is required for inaccurate flags"}), 400

    user_email = _current_user_email()
    status = "positive" if flag_type == "accurate" else "pending"

    conn = None
    try:
        conn = get_db_conn()
        cur  = conn.cursor()

        # Enforce one rating per user per search_event_id
        if search_event_id:
            cur.execute(
                "SELECT id FROM flagged_queries WHERE search_event_id = %s AND flagged_by = %s LIMIT 1",
                (search_event_id, user_email),
            )
            existing = cur.fetchone()
            if existing:
                # Update the existing rating instead of inserting a duplicate
                cur.execute(
                    "UPDATE flagged_queries SET flag_type=%s, container=%s, explanation=%s, "
                    "request_article=%s, section_ratings=%s, status=%s, flagged_at=NOW() "
                    "WHERE id=%s",
                    (flag_type, container, explanation, request_article,
                     json.dumps(section_ratings), status, existing[0]),
                )
                conn.commit()
                return jsonify({"success": True, "flag_id": existing[0], "updated": True})

        flag_id = str(uuid.uuid4())
        cur.execute(
            "INSERT INTO flagged_queries "
            "(id, query, loop_type, answer, explanation, flagged_by, status, "
            " flag_type, container, request_article, "
            " branch, confidence_level, duration_sec, "
            " input_tokens, output_tokens, total_tokens, "
            " files_read, iterations, searches, "
            " search_event_id, section_ratings) "
            "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
            (flag_id, query, loop_type, answer, explanation, user_email, status,
             flag_type, container, request_article,
             data.get("branch", ""), data.get("confidence_level", ""),
             data.get("duration_sec", 0),
             data.get("input_tokens", 0), data.get("output_tokens", 0),
             data.get("total_tokens", 0),
             data.get("files_read", 0), data.get("iterations", 0),
             data.get("searches", 0),
             search_event_id, json.dumps(section_ratings)),
        )
        conn.commit()
        return jsonify({"success": True, "flag_id": flag_id})
    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/flag/check/<int:event_id>", methods=["GET"])
@login_required
def api_flag_check(event_id):
    """Check if the current user has already rated a specific search event."""
    user_email = _current_user_email()
    conn = None
    try:
        conn = get_db_conn()
        cur  = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(
            "SELECT id, flag_type, section_ratings, explanation, request_article "
            "FROM flagged_queries WHERE search_event_id = %s AND flagged_by = %s LIMIT 1",
            (event_id, user_email),
        )
        row = cur.fetchone()
        if row:
            return jsonify({"rated": True, "rating": dict(row)})
        return jsonify({"rated": False})
    except Exception as e:
        return jsonify({"rated": False, "error": str(e)})
    finally:
        if conn:
            put_db_conn(conn)


# ---------------------------------------------------------------------------
# Admin routes
# ---------------------------------------------------------------------------

ADMIN_TABLES = {
    "search_events": (
        "SELECT id, branch, confidence_level, duration_sec, error_message, "
        "files_read, input_tokens, iterations, output_tokens, "
        "query, ran_at, searches, status, total_tokens, "
        "COALESCE(model, '') AS model, "
        "COALESCE(user_email, '') AS user_email, "
        "COALESCE(response_sections, '') AS response_sections, "
        "COALESCE(job_id, '') AS job_id, "
        "COALESCE(screenshot_context, '') AS screenshot_context, "
        "(COALESCE(answer_text, '') != '') AS has_answer, "
        "(COALESCE(screenshot_b64, '') != '') AS has_screenshot "
        "FROM search_events ORDER BY ran_at DESC LIMIT 500"
    ),
    "flagged_queries": (
        "SELECT id, query, loop_type, "
        "COALESCE(answer, '') AS answer, "
        "COALESCE(explanation, '') AS explanation, "
        "COALESCE(flagged_by, '') AS flagged_by, "
        "flagged_at, status, "
        "COALESCE(reviewed_by, '') AS reviewed_by, "
        "reviewed_at, "
        "COALESCE(admin_notes, '') AS admin_notes, "
        "flag_type, "
        "COALESCE(container, '') AS container, "
        "request_article, "
        "COALESCE(branch, '') AS branch, "
        "COALESCE(confidence_level, '') AS confidence_level, "
        "duration_sec, input_tokens, output_tokens, total_tokens, "
        "files_read, iterations, searches "
        "FROM flagged_queries ORDER BY flagged_at DESC LIMIT 500"
    ),
    "api_jobs": (
        "SELECT id, status, query, "
        "COALESCE(branch, '') AS branch, "
        "COALESCE(user_email, 'external_api') AS user_email, "
        "COALESCE(error, '') AS error, "
        "created_at, completed_at "
        "FROM api_jobs ORDER BY created_at DESC LIMIT 500"
    ),
}


# /admin route removed — admin dashboard lives in the main page (toggle button)


@app.route("/admin/data/<table_name>")
@admin_required
def admin_table_data(table_name):
    if table_name not in ADMIN_TABLES:
        return jsonify({"error": "Invalid table name"}), 400
    rows = query_db(ADMIN_TABLES[table_name])
    return jsonify({"rows": rows, "count": len(rows)})


@app.route("/admin/export/<table_name>")
@admin_required
def admin_export_csv(table_name):
    if table_name not in ADMIN_TABLES:
        return "Invalid table name", 400
    rows = query_db(ADMIN_TABLES[table_name])
    if not rows:
        return "No data", 200
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=rows[0].keys())
    writer.writeheader()
    writer.writerows(rows)
    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": f"attachment; filename={table_name}.csv"},
    )


# ---------------------------------------------------------------------------
# Admin flag management routes
# ---------------------------------------------------------------------------

@app.route("/api/admin/stats")
@admin_required
def api_admin_stats():
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT COUNT(*) as total FROM flagged_queries")
        total = cur.fetchone()["total"]
        cur.execute("SELECT status, COUNT(*) as cnt FROM flagged_queries GROUP BY status")
        by_status = {row["status"]: row["cnt"] for row in cur.fetchall()}
        cur.execute(
            "SELECT query, COUNT(*) as cnt FROM flagged_queries "
            "GROUP BY query ORDER BY cnt DESC LIMIT 10"
        )
        top_queries = [{"query": row["query"], "count": row["cnt"]} for row in cur.fetchall()]
        cur.execute(
            "SELECT DATE(flagged_at) as day, COUNT(*) as cnt FROM flagged_queries "
            "WHERE flagged_at >= NOW() - INTERVAL '30 days' GROUP BY day ORDER BY day"
        )
        daily_trend = [{"date": str(row["day"]), "count": row["cnt"]} for row in cur.fetchall()]
        cur.execute(
            "SELECT loop_type, COUNT(*) as cnt FROM flagged_queries GROUP BY loop_type"
        )
        by_loop = {row["loop_type"]: row["cnt"] for row in cur.fetchall()}
        return jsonify({
            "total": total,
            "pending": by_status.get("pending", 0),
            "reviewed": by_status.get("reviewed", 0),
            "dismissed": by_status.get("dismissed", 0),
            "article_created": by_status.get("article_created", 0),
            "top_queries": top_queries,
            "daily_trend": daily_trend,
            "by_loop_type": by_loop,
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/admin/flags")
@admin_required
def api_admin_flags():
    status_filter = request.args.get("status", "")
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if status_filter:
            cur.execute(
                "SELECT * FROM flagged_queries WHERE status = %s ORDER BY flagged_at DESC",
                (status_filter,),
            )
        else:
            cur.execute("SELECT * FROM flagged_queries ORDER BY flagged_at DESC")
        rows = [dict(r) for r in cur.fetchall()]
        return jsonify({"flags": rows})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/admin/user-rankings")
@admin_required
def api_admin_user_rankings():
    """Return per-user aggregates: query count, token usage, estimated cost.

    Optional query params: start_date, end_date (YYYY-MM-DD) to scope the window.
    """
    conn = None
    try:
        start_date = request.args.get("start_date", "").strip()
        end_date   = request.args.get("end_date", "").strip()
        conn = get_db_conn()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        where_clause = ""
        params: list = []
        if start_date and end_date:
            where_clause = "WHERE ran_at::date >= %s::date AND ran_at::date <= %s::date"
            params = [start_date, end_date]
        cur.execute(f"""
            SELECT
                COALESCE(NULLIF(user_email, ''), 'unknown') AS user_email,
                COUNT(*)                                    AS query_count,
                SUM(COALESCE(input_tokens, 0))              AS total_input_tokens,
                SUM(COALESCE(output_tokens, 0))             AS total_output_tokens,
                SUM(COALESCE(total_tokens, 0))              AS total_tokens,
                COALESCE(NULLIF(model, ''), '{MODEL}')      AS model
            FROM search_events
            {where_clause}
            GROUP BY COALESCE(NULLIF(user_email, ''), 'unknown'),
                     COALESCE(NULLIF(model, ''), '{MODEL}')
            ORDER BY query_count DESC
        """, params if params else None)
        # Aggregate per-user across models
        user_agg: dict = {}
        for row in cur.fetchall():
            email = row["user_email"]
            inp = row["total_input_tokens"] or 0
            out = row["total_output_tokens"] or 0
            ci, co = _model_costs(row["model"])
            cost = (inp / 1_000_000) * ci + (out / 1_000_000) * co
            if email not in user_agg:
                user_agg[email] = {
                    "user_email": email, "query_count": 0,
                    "total_input_tokens": 0, "total_output_tokens": 0,
                    "total_tokens": 0, "estimated_cost": 0.0,
                }
            agg = user_agg[email]
            agg["query_count"]         += row["query_count"]
            agg["total_input_tokens"]  += inp
            agg["total_output_tokens"] += out
            agg["total_tokens"]        += row["total_tokens"] or 0
            agg["estimated_cost"]      += cost
        rankings = sorted(user_agg.values(), key=lambda r: r["query_count"], reverse=True)
        for r in rankings:
            r["estimated_cost"] = round(r["estimated_cost"], 4)
        return jsonify({"rankings": rankings})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/admin/flags/<flag_id>", methods=["PATCH"])
@admin_required
def api_admin_update_flag(flag_id):
    data = request.get_json() or {}
    new_status  = data.get("status", "")
    admin_notes = data.get("admin_notes", "")
    if new_status not in ("pending", "reviewed", "dismissed", "article_created"):
        return jsonify({"error": "Invalid status"}), 400
    conn = None
    try:
        conn = get_db_conn()
        cur  = conn.cursor()
        cur.execute(
            "UPDATE flagged_queries SET status=%s, admin_notes=%s, reviewed_by=%s, reviewed_at=NOW() WHERE id=%s",
            (new_status, admin_notes, _current_user_email(), flag_id),
        )
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


# ---------------------------------------------------------------------------
# Admin — individual record lookups
# ---------------------------------------------------------------------------

@app.route("/api/admin/search-events", methods=["GET"])
@admin_required
def api_admin_search_event_ids():
    """Return a list of all search event IDs with basic metadata for browsing."""
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(
            "SELECT id, ran_at, loop_type, query, branch, status, "
            "COALESCE(user_email, '') AS user_email, "
            "COALESCE(job_id, '') AS job_id, "
            "COALESCE(screenshot_context, '') AS screenshot_context "
            "FROM search_events ORDER BY ran_at DESC"
        )
        rows = [dict(r) for r in cur.fetchall()]
        return jsonify({
            "success": True,
            "count": len(rows),
            "events": rows,
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/admin/search-events/<int:event_id>", methods=["GET"])
@admin_required
def api_admin_search_event_detail(event_id):
    """Return full details for a single search event by ID."""
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT * FROM search_events WHERE id = %s", (event_id,))
        row = cur.fetchone()
        if not row:
            return jsonify({"error": f"Search event {event_id} not found"}), 404
        return jsonify({"success": True, "event": dict(row)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/admin/flags/<flag_id>/detail", methods=["GET"])
@admin_required
def api_admin_flag_detail(flag_id):
    """Return full details for a single flagged query by ID,
    including the original answer and all response containers."""
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT * FROM flagged_queries WHERE id = %s", (flag_id,))
        row = cur.fetchone()
        if not row:
            return jsonify({"error": f"Flagged query {flag_id} not found"}), 404
        return jsonify({"success": True, "flag": dict(row)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


# ---------------------------------------------------------------------------
# API Keys Management
# ---------------------------------------------------------------------------

@app.route("/api/admin/api-keys", methods=["POST"])
@admin_required
def admin_create_api_key():
    """Create a new API key. Returns the raw key once."""
    data = request.get_json() or {}
    label = data.get("label", "").strip()
    if not label:
        return jsonify({"error": "Label is required"}), 400

    # Generate key: "cora_" prefix + 48 random hex chars
    raw_key = "cora_" + secrets.token_hex(24)
    key_hash = hashlib.sha256(raw_key.encode()).hexdigest()
    key_prefix = raw_key[:12] + "..."
    key_id = str(uuid.uuid4())
    created_by = session.get("user", {}).get("email", "unknown")

    # Insert into DB
    conn = None
    try:
        conn = get_db_conn()
        conn.cursor().execute(
            "INSERT INTO api_keys (id, key_hash, key_prefix, label, created_by) VALUES (%s, %s, %s, %s, %s)",
            (key_id, key_hash, key_prefix, label, created_by),
        )
        conn.commit()
    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)

    # Return the raw key ONCE — it can never be retrieved again
    return jsonify({
        "success": True,
        "key_id": key_id,
        "api_key": raw_key,
        "key_prefix": key_prefix,
        "label": label,
        "message": "Save this key now — it cannot be retrieved again.",
    }), 201


@app.route("/api/admin/api-keys", methods=["GET"])
@admin_required
def admin_list_api_keys():
    """List all API keys (never exposes raw key or hash)."""
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(
            "SELECT id, key_prefix, label, created_by, created_at, last_used, is_active "
            "FROM api_keys ORDER BY created_at DESC"
        )
        keys = [dict(r) for r in cur.fetchall()]
        return jsonify({"success": True, "keys": keys})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/admin/api-keys/<key_id>", methods=["PATCH"])
@admin_required
def admin_update_api_key(key_id):
    """Revoke or reactivate a key, or update its label."""
    data = request.get_json() or {}
    is_active = data.get("is_active")
    label = data.get("label")

    updates = []
    params = []
    if is_active is not None:
        updates.append("is_active = %s")
        params.append(bool(is_active))
    if label is not None:
        updates.append("label = %s")
        params.append(label.strip())

    if not updates:
        return jsonify({"error": "No fields to update"}), 400

    params.append(key_id)
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute(f"UPDATE api_keys SET {', '.join(updates)} WHERE id = %s", params)
        if cur.rowcount == 0:
            return jsonify({"error": "Key not found"}), 404
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/admin/api-keys/<key_id>", methods=["DELETE"])
@admin_required
def admin_delete_api_key(key_id):
    """Permanently delete an API key."""
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("DELETE FROM api_keys WHERE id = %s", (key_id,))
        if cur.rowcount == 0:
            return jsonify({"error": "Key not found"}), 404
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


# ---------------------------------------------------------------------------
# Admin: Codebase Index management
# ---------------------------------------------------------------------------

@app.route("/api/admin/codebase-index")
@admin_required
def api_admin_codebase_index():
    """Return the full codebase index for the admin dashboard."""
    if not DATABASE_URL:
        return jsonify({"error": "No database configured"}), 500
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("""
            SELECT id, file_path, file_name, file_role, feature_area,
                   keywords, related_files, description, branch, hit_count,
                   discovered_at, last_verified, source_query,
                   module, layer, technology, ado_url
            FROM codebase_index
            ORDER BY hit_count DESC, last_verified DESC
        """)
        rows = [dict(r) for r in cur.fetchall()]
        # Summary stats
        cur.execute("""
            SELECT COUNT(*) AS total_files,
                   COUNT(DISTINCT feature_area) AS feature_areas,
                   COUNT(DISTINCT branch) AS branches,
                   SUM(hit_count) AS total_hits,
                   MAX(last_verified) AS last_updated
            FROM codebase_index
        """)
        stats = dict(cur.fetchone())
        return jsonify({"index": rows, "stats": stats})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/admin/codebase-index/<int:entry_id>", methods=["DELETE"])
@admin_required
def api_admin_delete_index_entry(entry_id):
    """Delete a single entry from the codebase index."""
    if not DATABASE_URL:
        return jsonify({"error": "No database configured"}), 500
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("DELETE FROM codebase_index WHERE id = %s", (entry_id,))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/admin/codebase-index/clear", methods=["POST"])
@admin_required
def api_admin_clear_index():
    """Clear the entire codebase index (reset learning)."""
    if not DATABASE_URL:
        return jsonify({"error": "No database configured"}), 500
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("DELETE FROM codebase_index")
        deleted = cur.rowcount
        conn.commit()
        return jsonify({"success": True, "deleted": deleted})
    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


@app.route("/api/admin/codebase-index/bulk", methods=["POST"])
@admin_required
def api_admin_bulk_import_index():
    """Bulk import codebase index entries.

    Expects JSON body:
    {
        "entries": [
            {
                "file_name": "FinancialMgmt",
                "file_path": "/code/.../FinancialMgmt",
                "file_role": "Financial Calculations",
                "feature_area": "Financial Management",
                "branch": "supported_release-1.25.2",
                "keywords": "EAC, ETC, budget...",
                "description": "Core financial management...",
                "related_files": "",
                "source_query": "pre-indexed"
            }, ...
        ],
        "skip_duplicates": true
    }
    """
    if not DATABASE_URL:
        return jsonify({"error": "No database configured"}), 500

    data = request.get_json()
    if not data or "entries" not in data:
        return jsonify({"error": "Missing 'entries' array in request body"}), 400

    entries = data["entries"]
    skip_dupes = data.get("skip_duplicates", True)
    upsert_mode = data.get("upsert", False)  # If true, update existing entries with new data

    if not isinstance(entries, list):
        return jsonify({"error": "'entries' must be an array"}), 400
    if len(entries) > 1000:
        return jsonify({"error": "Maximum 1000 entries per request"}), 400

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        inserted = 0
        updated = 0
        skipped = 0
        errors = 0

        for entry in entries:
            try:
                fp = entry.get("file_path", "")
                fn = entry.get("file_name", "")
                br = entry.get("branch", "supported_release-1.25.2")
                if not fn and not fp:
                    errors += 1
                    continue

                # Check if entry already exists
                cur.execute(
                    "SELECT id FROM codebase_index WHERE file_path = %s AND branch = %s",
                    (fp, br),
                )
                existing = cur.fetchone()

                if existing:
                    if upsert_mode:
                        # Update existing entry with enriched data (only non-empty fields)
                        cur.execute("""
                            UPDATE codebase_index SET
                                file_role    = CASE WHEN %(file_role)s != '' THEN %(file_role)s ELSE file_role END,
                                feature_area = CASE WHEN %(feature_area)s != '' THEN %(feature_area)s ELSE feature_area END,
                                keywords     = CASE WHEN %(keywords)s != '' THEN %(keywords)s ELSE keywords END,
                                description  = CASE WHEN %(description)s != '' THEN %(description)s ELSE description END,
                                related_files = CASE WHEN %(related_files)s != '' THEN %(related_files)s ELSE related_files END,
                                module       = CASE WHEN %(module)s != '' THEN %(module)s ELSE module END,
                                layer        = CASE WHEN %(layer)s != '' THEN %(layer)s ELSE layer END,
                                technology   = CASE WHEN %(technology)s != '' THEN %(technology)s ELSE technology END,
                                ado_url      = CASE WHEN %(ado_url)s != '' THEN %(ado_url)s ELSE ado_url END,
                                last_verified = NOW()
                            WHERE file_path = %(fp)s AND branch = %(br)s
                        """, {
                            "file_role": entry.get("file_role", ""),
                            "feature_area": entry.get("feature_area", ""),
                            "keywords": entry.get("keywords", ""),
                            "description": entry.get("description", ""),
                            "related_files": entry.get("related_files", ""),
                            "module": entry.get("module", ""),
                            "layer": entry.get("layer", ""),
                            "technology": entry.get("technology", ""),
                            "ado_url": entry.get("ado_url", ""),
                            "fp": fp, "br": br,
                        })
                        updated += 1
                    elif skip_dupes:
                        skipped += 1
                    continue

                cur.execute("""
                    INSERT INTO codebase_index
                        (file_name, file_path, file_role, feature_area, branch,
                         keywords, description, related_files, source_query, hit_count,
                         module, layer, technology, ado_url)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    fn, fp,
                    entry.get("file_role", ""),
                    entry.get("feature_area", ""),
                    br,
                    entry.get("keywords", ""),
                    entry.get("description", ""),
                    entry.get("related_files", ""),
                    entry.get("source_query", "pre-indexed"),
                    entry.get("hit_count", 0),
                    entry.get("module", ""),
                    entry.get("layer", ""),
                    entry.get("technology", ""),
                    entry.get("ado_url", ""),
                ))
                inserted += 1
            except Exception as row_err:
                print(f"[BulkImport] Row error for {entry.get('file_name', '?')}: {row_err}")
                errors += 1
                continue

        conn.commit()
        return jsonify({
            "success": True,
            "inserted": inserted,
            "updated": updated,
            "skipped": skipped,
            "errors": errors,
            "message": f"Bulk import complete. {inserted} inserted, {updated} updated, {skipped} skipped, {errors} errors.",
        }), 201
    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


# ---------------------------------------------------------------------------
# Admin: auto-index a branch from ADO file tree
# ---------------------------------------------------------------------------

def _categorize_ado_files(file_paths: list, branch: str) -> list:
    """Categorize a list of ADO file paths into codebase index entries.

    Returns only key architectural files: Controllers, Services, Repositories,
    Models, Helpers, Handlers, Middleware, key UI pages + CodeBehind,
    TypeScript, Telerik Reports, Stored Procedures, SQL Views/Functions,
    and build/deployment scripts.
    """
    ado_url_base = f"https://dev.azure.com/{ADO_ORG}/{ADO_PROJECT}/_git/{ADO_REPO}"
    key_pages = {'default','main','dashboard','list','detail','edit','manage','admin',
                 'settings','report','forecast','resource','project','task','schedule',
                 'budget','cost','risk','timesheet','document','import','export','login','search'}

    def _detect_area(p):
        pl = p.lower()
        checks = [('forecast','Financial Management'),('resource','Resource Management'),
                   ('risk','Risk Management'),('budget','Financial Management'),
                   ('cost','Financial Management'),('financ','Financial Management'),
                   ('schedule','Task Management'),('gantt','Task Management'),
                   ('report','Reporting'),('admin','Administration'),
                   ('configuration','Administration'),('settings','Administration'),
                   ('import','Data Import/Export'),('export','Data Import/Export'),
                   ('dashboard','Reporting'),('task','Task Management'),
                   ('activity','Task Management'),('deliverable','Task Management'),
                   ('project','Project Management'),('login','Administration'),
                   ('auth','Administration'),('account','Administration'),
                   ('calendar','Calendar'),('document','Content Management'),
                   ('attachment','Content Management'),('notification','Communication'),
                   ('email','Communication'),('audit','Audit'),('history','Audit'),
                   ('custom','Extensibility'),('smartform','Extensibility'),
                   ('udf','Extensibility'),('timesheet','Time Tracking'),
                   ('timeentry','Time Tracking'),('material','Materials'),
                   ('integration','Integration'),('webservice','Integration')]
        for k, v in checks:
            if k in pl:
                return v
        return 'Project Management'

    def _detect_module(p):
        checks = [('PVPortal','PVPortal (Aurelia SPA)'),('ProjectVisionLib','ProjectVisionLib'),
                   ('CORA.Common','CORA.Common'),('AccountLibrary','AccountLibrary'),
                   ('CoraCustomControls','CoraCustomControls'),('ImportExport','ImportExport'),
                   ('JobsScheduler','Jobs/Scheduler'),('BizScheduler','Jobs/Scheduler'),
                   ('HPMS','HPMS'),('EACAutomation','EACAutomation'),
                   ('LicenseKey','LicenseKey'),('SyncfusionHelper','SyncfusionHelper'),
                   ('/scripts/','DevOps Scripts'),('/reports/','Telerik Reports'),
                   ('coraportal','Cora Portal'),('coraapi','Cora API'),
                   ('projectvision','ProjectVision Web')]
        for k, v in checks:
            if k in p:
                return v
        return ''

    skip_dirs = {'/binaries/', '/node_modules/', '/packages/', '/obj/', '/bin/'}
    entries = []

    for fp in file_paths:
        name = fp.rsplit('/', 1)[-1] if '/' in fp else fp
        nl = name.lower()
        ext_pos = nl.rfind('.')
        ext = nl[ext_pos:] if ext_pos >= 0 else ''
        pl = fp.lower()

        if any(d in pl for d in skip_dirs):
            continue

        role = layer = tech = ''

        if ext == '.cs':
            if 'controller' in nl:          role, layer, tech = 'Controller', 'API', 'C#/.NET'
            elif 'service' in nl and '.designer.' not in nl:
                                            role, layer, tech = 'Service', 'Business Logic', 'C#/.NET'
            elif 'repository' in nl:        role, layer, tech = 'Repository', 'Data Access', 'C#/.NET'
            elif 'dto' in nl or nl.endswith('model.cs') or 'viewmodel' in nl:
                                            role, layer, tech = 'Model', 'Application', 'C#/.NET'
            elif any(k in nl for k in ('helper','util','extension')):
                                            role, layer, tech = 'Helper', 'Application', 'C#/.NET'
            elif 'handler' in nl:           role, layer, tech = 'Handler', 'Application', 'C#/.NET'
            elif 'middleware' in nl:         role, layer, tech = 'Middleware', 'Application', 'C#/.NET'

        elif ext == '.vb':
            if (nl.endswith('.aspx.vb') or nl.endswith('.ascx.vb')) and any(k in nl for k in key_pages):
                                            role, layer, tech = 'CodeBehind', 'Presentation', 'VB.NET'
            elif 'service' in nl:           role, layer, tech = 'Service', 'Business Logic', 'VB.NET'
            elif 'repository' in nl or 'dataaccess' in nl:
                                            role, layer, tech = 'Repository', 'Data Access', 'VB.NET'
            elif 'helper' in nl or 'util' in nl:
                                            role, layer, tech = 'Helper', 'Application', 'VB.NET'

        elif ext == '.aspx' and any(k in nl for k in key_pages):
            role, layer, tech = 'UI', 'Presentation', 'ASP.NET WebForms'

        elif ext == '.ts' and not nl.endswith('.d.ts'):
            role, layer, tech = 'TypeScript', 'Frontend', 'TypeScript/Aurelia'

        elif ext == '.trdx':
            role, layer, tech = 'Telerik Report', 'Reporting', 'Telerik TRDX'

        elif ext == '.sql':
            if any(nl.startswith(p) for p in ('sp_','usp_')) or 'proc' in nl:
                role, layer, tech = 'Stored Procedure', 'Database', 'SQL Server'
            elif nl.startswith('vw_') or 'view' in nl:
                role, layer, tech = 'SQL View', 'Database', 'SQL Server'
            elif any(nl.startswith(p) for p in ('fn_','uf_')) or 'function' in nl:
                role, layer, tech = 'SQL Function', 'Database', 'SQL Server'

        elif fp.startswith('/scripts/') and ext in ('.ps1','.bat','.yaml','.yml'):
            role = 'Deployment Script' if 'deployment' in fp else ('Build Script' if 'builds' in fp else 'Script')
            layer, tech = 'Infrastructure', ('PowerShell' if ext == '.ps1' else ('YAML/Azure Pipelines' if ext in ('.yaml','.yml') else 'Batch'))

        elif fp == '/BuildAndBuildValidation.yaml':
            role, layer, tech = 'CI/CD Pipeline', 'Infrastructure', 'YAML/Azure Pipelines'

        if not role:
            continue

        from urllib.parse import quote
        ado_url = f"{ado_url_base}?path={quote(fp)}&version=GB{branch}"
        desc = ''
        kw = ''
        if ext == '.trdx':
            parts = fp.split('/')
            desc = (parts[2] if len(parts) > 2 else '') + ' report'
            kw = 'report,telerik,dashboard,analytics,metrics'

        entries.append({
            "file_path": fp, "file_name": name, "file_role": role,
            "feature_area": _detect_area(fp), "module": _detect_module(fp),
            "layer": layer, "technology": tech, "description": desc,
            "keywords": kw, "branch": branch,
            "ado_url": ado_url, "hit_count": 0,
        })

    return entries


@app.route("/api/admin/codebase-index/auto-index", methods=["POST"])
@admin_required
def api_admin_auto_index_branch():
    """Discover files on a branch via the ADO Items API and index key architectural files.

    POST JSON: {"branch": "supported_release-1.25.1"}
    Optional:  {"branch": "...", "clear_existing": true}  to wipe existing entries for that branch first.
    """
    import requests as _requests
    from urllib.parse import quote as _quote

    data = request.get_json(force=True)
    branch = data.get("branch", "").strip()
    if not branch:
        return jsonify({"error": "branch is required"}), 400

    clear_existing = data.get("clear_existing", False)
    pat = os.environ.get("AZURE_DEVOPS_PAT", "")
    if not pat:
        return jsonify({"error": "AZURE_DEVOPS_PAT not configured"}), 500

    ado_token = base64.b64encode(f":{pat}".encode()).decode()
    headers = {"Authorization": f"Basic {ado_token}", "Content-Type": "application/json"}

    # Fetch file tree from key directories
    items_base = f"https://dev.azure.com/{ADO_ORG}/{ADO_PROJECT}/_apis/git/repositories/{ADO_REPO}/items"
    v_param = f"&versionDescriptor.version={_quote(branch)}&versionDescriptor.versionType=branch&$format=json&api-version=7.0"
    scope_dirs = [
        "/code/ProjectvisionSolution_2008", "/code/PVPortal",
        "/code/inetpub", "/code/PPM.SyncfusionHelper",
        "/scripts", "/reports",
    ]

    all_paths = []
    fetch_errors = []
    for scope in scope_dirs:
        url = f"{items_base}?scopePath={_quote(scope)}&recursionLevel=Full&includeContentMetadata=true{v_param}"
        try:
            resp = _requests.get(url, headers=headers, timeout=120)
            if resp.ok:
                items = resp.json().get("value", [])
                all_paths.extend(item["path"] for item in items if not item.get("isFolder"))
            else:
                fetch_errors.append(f"{scope}: HTTP {resp.status_code}")
        except Exception as e:
            fetch_errors.append(f"{scope}: {e}")

    # Root-level files
    url = f"{items_base}?scopePath=/&recursionLevel=OneLevel&includeContentMetadata=true{v_param}"
    try:
        resp = _requests.get(url, headers=headers, timeout=30)
        if resp.ok:
            items = resp.json().get("value", [])
            all_paths.extend(item["path"] for item in items if not item.get("isFolder"))
    except Exception as e:
        fetch_errors.append(f"/: {e}")

    if not all_paths:
        return jsonify({"error": "No files found", "fetch_errors": fetch_errors}), 404

    # Categorize
    entries = _categorize_ado_files(all_paths, branch)

    if not entries:
        return jsonify({"error": "No key architectural files found", "total_files": len(all_paths)}), 404

    # Import into database
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()

        if clear_existing:
            cur.execute("DELETE FROM codebase_index WHERE branch = %s", (branch,))
            cleared = cur.rowcount
        else:
            cleared = 0

        inserted = 0
        skipped = 0
        for e in entries:
            cur.execute(
                "SELECT id FROM codebase_index WHERE file_path = %s AND branch = %s",
                (e["file_path"], branch),
            )
            if cur.fetchone():
                skipped += 1
                continue
            cur.execute("""
                INSERT INTO codebase_index
                    (file_name, file_path, file_role, feature_area, branch,
                     keywords, description, related_files, source_query, hit_count,
                     module, layer, technology, ado_url)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                e["file_name"], e["file_path"], e["file_role"], e["feature_area"],
                branch, e["keywords"], e["description"], "", "auto-indexed", e["hit_count"],
                e["module"], e["layer"], e["technology"], e["ado_url"],
            ))
            inserted += 1

        conn.commit()
        return jsonify({
            "success": True,
            "branch": branch,
            "total_files_scanned": len(all_paths),
            "entries_categorized": len(entries),
            "inserted": inserted,
            "skipped_existing": skipped,
            "cleared_existing": cleared,
            "fetch_errors": fetch_errors,
        })

    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)


# ---------------------------------------------------------------------------
# Auth routes — Entra ID OAuth
# ---------------------------------------------------------------------------

@app.route("/login")
def login():
    if AUTH_DISABLED:
        return redirect(url_for("index"))
    msal_app = _build_msal_app()
    state = secrets.token_urlsafe(16)
    session["auth_state"] = state
    auth_url = msal_app.get_authorization_request_url(
        ENTRA_SCOPES,
        redirect_uri=ENTRA_REDIRECT_URI,
        state=state,
    )
    return redirect(auth_url)


@app.route("/auth/callback")
def auth_callback():
    if AUTH_DISABLED:
        return redirect(url_for("index"))
    if request.args.get("state") != session.get("auth_state"):
        return "State mismatch — possible CSRF attack.", 400
    code = request.args.get("code")
    if not code:
        return f"Auth error: {request.args.get('error_description', 'No code returned')}", 400
    msal_app = _build_msal_app()
    result = msal_app.acquire_token_by_authorization_code(
        code,
        scopes=ENTRA_SCOPES,
        redirect_uri=ENTRA_REDIRECT_URI,
    )
    if "error" in result:
        return f"Auth error: {result.get('error_description', result['error'])}", 400
    # Fetch user profile from Microsoft Graph
    graph_resp = requests.get(
        "https://graph.microsoft.com/v1.0/me",
        headers={"Authorization": f"Bearer {result['access_token']}"},
    )
    if graph_resp.status_code == 200:
        session["user"] = graph_resp.json()
    else:
        session["user"] = {"displayName": "User", "mail": "unknown"}
    return redirect(url_for("index"))


@app.route("/logout")
def logout():
    session.clear()
    if AUTH_DISABLED:
        return redirect(url_for("index"))
    return redirect(
        f"{ENTRA_AUTHORITY}/oauth2/v2.0/logout"
        f"?post_logout_redirect_uri={urllib.parse.quote(ENTRA_REDIRECT_URI.rsplit('/auth/callback', 1)[0], safe='')}"
    )


# ---------------------------------------------------------------------------
# External API — unified query endpoint for other apps (Jira, n8n, etc.)
# ---------------------------------------------------------------------------
EXTERNAL_API_KEY = os.environ.get("EXTERNAL_API_KEY", "")


def _check_api_key():
    """Validate bearer token for external API calls. Returns error response or None.
    Checks database keys first, then falls back to EXTERNAL_API_KEY env var."""
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        return jsonify({"error": "Missing or invalid Authorization header. Use: Bearer <api_key>"}), 401

    token = auth_header[7:].strip()
    token_hash = hashlib.sha256(token.encode()).hexdigest()

    # Try to find the key in the database
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(
            "SELECT id, label FROM api_keys WHERE key_hash = %s AND is_active = TRUE",
            (token_hash,)
        )
        row = cur.fetchone()
        if row:
            # Found a valid DB key — update last_used and store info in g
            cur.execute(
                "UPDATE api_keys SET last_used = NOW() WHERE id = %s",
                (row['id'],)
            )
            conn.commit()
            g.api_key_id = row['id']
            g.api_key_label = row['label']
            return None
    except Exception as e:
        print(f"[API Key] Database check failed: {e}")
    finally:
        if conn:
            put_db_conn(conn)

    # Fall back to EXTERNAL_API_KEY env var
    if EXTERNAL_API_KEY and secrets.compare_digest(token, EXTERNAL_API_KEY):
        g.api_key_id = "env_default"
        g.api_key_label = "Environment Default"
        return None

    # No valid key found
    if not EXTERNAL_API_KEY and not token:
        return jsonify({"error": "External API not configured — set EXTERNAL_API_KEY env var or create an API key."}), 503

    return jsonify({"error": "Invalid API key."}), 403


# ---------------------------------------------------------------------------
# Async job runner — processes API queries in background threads
# ---------------------------------------------------------------------------

def _run_job_in_background(job_id, query, branch, user_email,
                           screenshot_context="", screenshot_b64="", screenshot_mime="",
                           model=""):
    """Execute the agentic search for a job and write results to the DB."""
    _model = model or MODEL
    t0 = time.time()
    try:
        result = run_agentic_loop(query, branch, model=_model)
        total_time = round(time.time() - t0, 1)

        # Track in search_events
        track_search_event(
            "external_api", query, branch, "success",
            total_time,
            result.get("total_files_read", 0),
            result.get("total_iterations", 0),
            result.get("total_searches", 0),
            result.get("confidence", {}).get("level", ""),
            input_tokens=result.get("input_tokens", 0),
            output_tokens=result.get("output_tokens", 0),
            user_email=user_email,
            answer=result.get("answer", ""),
            job_id=job_id,
            screenshot_context=screenshot_context,
            screenshot_b64=screenshot_b64,
            screenshot_mime=screenshot_mime,
            model=_model,
        )

        # Complete the api_jobs row (adds screenshot_context to payload)
        result["branch"] = branch
        _complete_api_job(job_id, result, screenshot_context=screenshot_context)

    except Exception as e:
        total_time = round(time.time() - t0, 1)
        _fail_api_job(job_id, str(e), total_time)


@app.route("/api/query", methods=["POST"])
def api_v1_query():
    """
    External query endpoint (async).

    Submits a query for processing and returns a job_id immediately.
    Poll /api/query/<job_id> for results.

    Optionally include a screenshot — Claude will extract UI context and
    append it to the query automatically, just like the web UI.

    Accepts two content types:

    1) application/json — screenshot as base64 in the JSON body:
        {
            "query":  "Why is the EAC Submit button greyed out?",
            "branch": "supported_release-1.25.2",
            "screenshot": {
                "image_b64": "<base64-encoded image>",
                "mime_type": "image/png"
            }
        }

    2) multipart/form-data — screenshot as a file upload:
        query:      "Why is the EAC Submit button greyed out?"
        branch:     "supported_release-1.25.2"    (optional)
        screenshot: <file>                         (optional)

    Response JSON (202 Accepted):
        {
            "job_id":    "abc-123",
            "status":    "processing",
            "poll_url":  "/api/query/abc-123",
            "message":   "Query submitted. Poll poll_url for results.",
            "screenshot_context": "..."
        }

    Auth: Bearer token via EXTERNAL_API_KEY env var.
    """
    # ── Auth ──
    auth_err = _check_api_key()
    if auth_err:
        return auth_err

    # ── Parse request (JSON or multipart/form-data) ──
    image_b64 = ""
    mime_type = "image/png"

    req_model = ""
    if request.content_type and "multipart/form-data" in request.content_type:
        # ── Multipart: fields + file upload ──
        query  = (request.form.get("query") or "").strip()
        branch = (request.form.get("branch") or "").strip()
        req_model = (request.form.get("model") or "").strip()
        file   = request.files.get("screenshot")
        if file and file.filename:
            image_b64 = base64.b64encode(file.read()).decode()
            mime_type  = file.content_type or "image/png"
    else:
        # ── JSON body ──
        data   = request.get_json(force=True) or {}
        query  = data.get("query", "").strip()
        branch = data.get("branch", "").strip()
        req_model = data.get("model", "").strip()
        screenshot_data = data.get("screenshot")
        if screenshot_data and isinstance(screenshot_data, dict):
            image_b64 = screenshot_data.get("image_b64", "").strip()
            mime_type  = screenshot_data.get("mime_type", "image/png").strip()

    _model = req_model if req_model in ALLOWED_MODELS else MODEL

    if not query:
        return jsonify({"error": "No query provided."}), 400

    # ── Resolve branch ──
    if not branch:
        try:
            branches = ado_get_branches()
            branch = branches[0] if branches else "main"
        except Exception:
            branch = "main"

    raw_email = request.headers.get("X-User-Email", "").strip()
    user_email = f"external_api ({raw_email})" if raw_email else "external_api"

    # ── Extract screenshot context (if provided) ──
    screenshot_context = ""
    if image_b64:
        try:
            screenshot_context = _extract_screenshot_context(image_b64, mime_type, model=_model)
        except Exception as e:
            return jsonify({"error": f"Screenshot extraction failed: {e}"}), 400

    # ── Build final query with screenshot context (matches UI behavior) ──
    final_query = query
    if screenshot_context:
        final_query += "\n\n[Screenshot context — extracted from attached UI screenshot]\n" + screenshot_context

    # ── Create job and return immediately ──
    job_id = str(uuid.uuid4())
    conn = None
    try:
        conn = get_db_conn()
        conn.cursor().execute(
            "INSERT INTO api_jobs (id, status, query, branch, user_email, model) VALUES (%s, 'processing', %s, %s, %s, %s)",
            (job_id, query, branch, user_email, _model),
        )
        conn.commit()
    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({"error": f"Could not create job: {e}"}), 500
    finally:
        if conn:
            put_db_conn(conn)

    t = threading.Thread(target=_run_job_in_background, args=(job_id, final_query, branch, user_email, screenshot_context, image_b64, mime_type, _model), daemon=True)
    t.start()

    response = {
        "job_id":   job_id,
        "status":   "processing",
        "poll_url": f"/api/query/{job_id}",
        "message":  "Query submitted. Poll poll_url for results.",
        "model":    _model,
    }
    if screenshot_context:
        response["screenshot_context"] = screenshot_context
    return jsonify(response), 202


@app.route("/api/query/<job_id>", methods=["GET"])
def api_v1_query_status(job_id):
    """
    Poll for async query results.

    Returns:
      - 200 with full result when job is completed or failed
      - 200 with status "processing" while the search is still running
      - 404 if job_id not found

    Query parameters for completed jobs:
      - fields: comma-separated list of top-level fields to return.
                Valid values: answer, sections, metadata.
                Default: all fields returned.
      - section: return only a specific section by name.
                 Valid values: overview, user_guide, technical_reference.
                 When set, returns only that section's content under a "section" key.

    Auth: Bearer token via EXTERNAL_API_KEY env var.
    """
    auth_err = _check_api_key()
    if auth_err:
        return auth_err

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT * FROM api_jobs WHERE id=%s", (job_id,))
        row = cur.fetchone()
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if conn:
            put_db_conn(conn)

    if not row:
        return jsonify({"error": "Job not found.", "job_id": job_id}), 404

    status = row["status"]

    _row_model = row.get("model") or MODEL

    if status == "processing":
        resp = {
            "job_id":  job_id,
            "status":  "processing",
            "message": "Search is still running. Poll again shortly.",
            "query":   row.get("query", ""),
            "model":   _row_model,
        }
        # If result_json already has screenshot_context (shouldn't for processing, but just in case)
        if row.get("result_json"):
            try:
                partial = json.loads(row["result_json"])
                if partial.get("screenshot_context"):
                    resp["screenshot_context"] = partial["screenshot_context"]
            except Exception:
                pass
        return jsonify(resp)

    # completed or failed — parse stored result
    result_json = row["result_json"]
    if result_json:
        result = json.loads(result_json)
    else:
        result = {"success": False, "error": row["error"] or "Unknown error"}

    # Always surface query and screenshot_context so callers know what was submitted
    _query = row.get("query", "")
    _sc    = result.get("screenshot_context", "")
    result.setdefault("query", _query)
    if _sc:
        result.setdefault("screenshot_context", _sc)

    # ── Apply response filters (only for successful completions) ──
    fields_param  = request.args.get("fields", "").strip()
    section_param = request.args.get("section", "").strip()

    VALID_SECTIONS = {"overview", "user_guide", "technical_reference"}

    # If a specific section is requested, return just that section
    if section_param and result.get("success"):
        if section_param not in VALID_SECTIONS:
            return jsonify({
                "error": f"Invalid section '{section_param}'. "
                         f"Valid values: {', '.join(sorted(VALID_SECTIONS))}",
            }), 400
        sections = result.get("sections", {})
        resp = {
            "job_id":       job_id,
            "status":       status,
            "success":      True,
            "query":        _query,
            "model":        _row_model,
            "section_name": section_param,
            "section":      sections.get(section_param, ""),
        }
        if _sc:
            resp["screenshot_context"] = _sc
        return jsonify(resp)

    # If specific fields are requested, filter the response
    if fields_param and result.get("success"):
        VALID_FIELDS = {"answer", "sections", "metadata"}
        requested = {f.strip() for f in fields_param.split(",") if f.strip()}
        invalid = requested - VALID_FIELDS
        if invalid:
            return jsonify({
                "error": f"Invalid fields: {', '.join(sorted(invalid))}. "
                         f"Valid values: {', '.join(sorted(VALID_FIELDS))}",
            }), 400
        filtered = {
            "job_id":  job_id,
            "status":  status,
            "success": True,
            "query":   _query,
            "model":   _row_model,
        }
        if _sc:
            filtered["screenshot_context"] = _sc
        for field in requested:
            if field in result:
                filtered[field] = result[field]
        filtered["created_at"]   = row["created_at"]
        filtered["completed_at"] = row["completed_at"]
        return jsonify(filtered)

    # Default: return everything
    result["job_id"]       = job_id
    result["status"]       = status
    result["model"]        = _row_model
    result["created_at"]   = row["created_at"]
    result["completed_at"] = row["completed_at"]
    return jsonify(result)


@app.route("/api/health", methods=["GET"])
def api_v1_health():
    """Public health check for the external API."""
    return jsonify({
        "status":    "ok",
        "api_ready": bool(EXTERNAL_API_KEY),
    })


if __name__ == "__main__":
    app.run(debug=True, port=5000)
